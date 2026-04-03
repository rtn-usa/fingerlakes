#!/usr/bin/env python3
"""Generate the fixed Cayuga Lake chapter as a properly formatted .docx"""

from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os

# --- Setup ---
doc = Document()

# Set default font to Arial
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Set page size US Letter and 1-inch margins
for section in doc.sections:
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- Define Hyperlink character style so links are clickable ---
styles_element = doc.styles.element
hyperlink_style = parse_xml(
    f'<w:style {nsdecls("w")} w:type="character" w:styleId="Hyperlink">'
    f'  <w:name w:val="Hyperlink"/>'
    f'  <w:basedOn w:val="DefaultParagraphFont"/>'
    f'  <w:uiPriority w:val="99"/>'
    f'  <w:unhideWhenUsed/>'
    f'  <w:rPr>'
    f'    <w:color w:val="0563C1"/>'
    f'    <w:u w:val="single"/>'
    f'  </w:rPr>'
    f'</w:style>'
)
styles_element.append(hyperlink_style)

# --- Hyperlink colors ---
WIKI_BLUE = '0563C1'
MAPS_GREEN = '1A7340'
UPSTATE_RED = 'C0392B'
OTHER_PURPLE = '6C3483'
CTA_BG = 'EAF4FB'

def add_hyperlink(paragraph, text, url, color_hex):
    """Add a colored, underlined hyperlink to a paragraph using OxmlElement."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), color_hex)
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')
    rPr.append(sz)

    run_elem.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)
    return paragraph

def add_run(paragraph, text, bold=False, italic=False, size=None):
    """Add a plain text run."""
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    return run

def add_heading_styled(text, level=1):
    """Add heading with Arial font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
    return h

def add_image_placeholder(caption):
    """Add a bracketed image placeholder."""
    p = doc.add_paragraph()
    run = p.add_run(f'[ {caption} ]')
    run.font.name = 'Arial'
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p

def add_image_caption(text):
    """Add an image caption in italic."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.italic = True
    run.font.size = Pt(10)
    return p

# ============================================================
# CHAPTER HEADER
# ============================================================
header_p = doc.add_paragraph()
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = header_p.add_run('FINGER LAKES / CAYUGA LAKE')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

h1 = add_heading_styled('Cayuga Lake', level=1)

tagline_p = doc.add_paragraph()
run = tagline_p.add_run('The longest lake. The deepest reform corridor. Thirty-nine miles that changed American history twice and is still producing Riesling that earns international attention.')
run.font.name = 'Arial'
run.italic = True
run.font.size = Pt(11)

meta_p = doc.add_paragraph()
run = meta_p.add_run('Chapter draft  \u00b7  Cayuga, Seneca, and Tompkins counties  \u00b7  March 2026')
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# IMAGE PLACEHOLDER
add_image_placeholder('IMAGE: Cayuga Lake aerial \u2014 length and valley visible, late summer or early fall preferred')
cap = doc.add_paragraph()
add_run(cap, 'Cayuga Lake', italic=True)
add_run(cap, ', looking south toward ', italic=True)
add_hyperlink(cap, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(cap, '. 38.9 miles, 435 feet deep, named for the people who farmed its shores for centuries.', italic=True)

# ============================================================
# TLDR - "The long and short of it"
# ============================================================
add_heading_styled('The long and short of it', level=3)
tldr_p = doc.add_paragraph()
add_run(tldr_p, 'Cayuga is the longest lake in the region and the one that rewards people who stay more than two nights. The reform corridor between ')
add_hyperlink(tldr_p, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(tldr_p, ' and ')
add_hyperlink(tldr_p, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(tldr_p, ' is the most historically consequential fifteen miles in New York State. Base in ')
add_hyperlink(tldr_p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(tldr_p, ' or ')
add_hyperlink(tldr_p, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(tldr_p, ', run the east shore on day two, and plan at least one morning at the Farmers Market before you do anything else.')

# ============================================================
# CHAPTER INTRO - "The lake, and what it keeps producing"
# ============================================================
add_heading_styled('The lake, and what it keeps producing', level=2)

# P1: Physical character
p1 = doc.add_paragraph()
add_run(p1, 'Cayuga Lake is 38.9 miles long and 435 feet deep, the longest of the eleven ')
add_hyperlink(p1, 'Finger Lakes', 'https://en.wikipedia.org/wiki/Finger_Lakes', WIKI_BLUE)
add_run(p1, ' and the second deepest. It runs north-south through three counties: ')
add_hyperlink(p1, 'Cayuga', 'https://en.wikipedia.org/wiki/Cayuga_County,_New_York', WIKI_BLUE)
add_run(p1, ', ')
add_hyperlink(p1, 'Seneca', 'https://en.wikipedia.org/wiki/Seneca_County,_New_York', WIKI_BLUE)
add_run(p1, ', and ')
add_hyperlink(p1, 'Tompkins', 'https://en.wikipedia.org/wiki/Tompkins_County,_New_York', WIKI_BLUE)
add_run(p1, '. ')
add_hyperlink(p1, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p1, ' sits at the southern tip where the land rises sharply on both sides and the gorges begin. ')
add_hyperlink(p1, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p1, ' anchors the north, where ')
add_hyperlink(p1, 'Harriet Tubman', 'https://en.wikipedia.org/wiki/Harriet_Tubman', WIKI_BLUE)
add_run(p1, ' lived for the last 54 years of her life. More than 150 waterfalls lie within ten miles of downtown ')
add_hyperlink(p1, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p1, ', and the glacial geology that carved the gorges is the same geology that makes the vineyards possible.')

# P2: Indigenous and early settlement history
p2 = doc.add_paragraph()
add_run(p2, 'The ')
add_hyperlink(p2, 'Cayuga Nation', 'https://en.wikipedia.org/wiki/Cayuga_Nation', WIKI_BLUE)
add_run(p2, ', Gayogo\u0331ho\u0301:n\u01eb\u02bc in their own language, are the People of the Great Swamp and one of the founding nations of the ')
add_hyperlink(p2, 'Haudenosaunee Confederacy', 'https://en.wikipedia.org/wiki/Iroquois', WIKI_BLUE)
add_run(p2, '. They held the territory around Cayuga Lake for centuries, sustained by the same agricultural conditions that make the land productive today. The ')
add_hyperlink(p2, 'Sullivan-Clinton Campaign', 'https://en.wikipedia.org/wiki/Sullivan%E2%80%93Clinton_campaign', WIKI_BLUE)
add_run(p2, ' burned those towns in 1779. The Cayuga were dispossessed under treaties negotiated in the 1790s that the Nation regards as illegitimate to this day. The lake still has their name.')

# P3: Wine/agricultural/industrial identity
p3 = doc.add_paragraph()
add_run(p3, 'The ')
add_hyperlink(p3, 'Cayuga Lake Wine Trail', 'https://upstate.tourismo.app/trails/cayuga-lake-wine-trail', UPSTATE_RED)
add_run(p3, ' was the first organized wine trail in the United States, established in 1983. There are 21 producers in the current passport. The farms along Route 89 have been there longer than the trail designation. The reform corridor between ')
add_hyperlink(p3, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p3, ' and ')
add_hyperlink(p3, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(p3, ' contains two national parks, the Underground Railroad\'s most consequential conductor, and the founding document of the American women\'s rights movement. That is Cayuga Lake\'s character, stated plainly: a place where the conditions have always attracted people who recognize what they\'re looking at.')

# P4: Connecting thread (Aurora story - original intro content preserved)
p4 = doc.add_paragraph()
add_hyperlink(p4, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(p4, ' is under one square mile. In 1833, a man named E.B. Morgan built an inn on Aurora\'s main street. Morgan was a co-founder of the ')
add_hyperlink(p4, 'New York Times', 'https://en.wikipedia.org/wiki/The_New_York_Times', WIKI_BLUE)
add_run(p4, '. In 1852, Henry Wells moved his family here and began building a house on a strip of land between the lake and a ravine, which he named Glen Park. Wells had recently co-founded two companies: ')
add_hyperlink(p4, 'American Express', 'https://en.wikipedia.org/wiki/American_Express', WIKI_BLUE)
add_run(p4, ' and ')
add_hyperlink(p4, 'Wells Fargo', 'https://en.wikipedia.org/wiki/Wells_Fargo', WIKI_BLUE)
add_run(p4, '. He considered higher education for women to be, in his words, "the dream of my life." In 1868 he founded ')
add_hyperlink(p4, 'Wells College', 'https://en.wikipedia.org/wiki/Wells_College', WIKI_BLUE)
add_run(p4, ' on the land adjacent to Glen Park. Future president ')
add_hyperlink(p4, 'Millard Fillmore', 'https://en.wikipedia.org/wiki/Millard_Fillmore', WIKI_BLUE)
add_run(p4, ' had already attended the Cayuga Lake Academy in ')
add_hyperlink(p4, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(p4, ' a generation earlier. In 1962, a young woman named Pleasant Rowland graduated from ')
add_hyperlink(p4, 'Wells College', 'https://en.wikipedia.org/wiki/Wells_College', WIKI_BLUE)
add_run(p4, '. She went on to create the ')
add_hyperlink(p4, 'American Girl', 'https://en.wikipedia.org/wiki/American_Girl', WIKI_BLUE)
add_run(p4, ' doll line, sold it to Mattel in 1998 for $700 million, and came back to ')
add_hyperlink(p4, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(p4, '. The ')
add_hyperlink(p4, 'Aurora Inn', 'https://www.innsofaurora.com/', OTHER_PURPLE)
add_run(p4, ', Morgan\'s building, had closed in 2000. The village was failing, incrementally and then all at once. Rowland bought the inn and restored it. Then she bought more buildings. Then she bought ')
add_hyperlink(p4, 'MacKenzie-Childs', 'https://www.mackenzie-childs.com/', OTHER_PURPLE)
add_run(p4, ', the home goods company headquartered on a farm just south of the village, when it was on the verge of closing. The ')
add_hyperlink(p4, 'Inns of Aurora', 'https://www.innsofaurora.com/', OTHER_PURPLE)
add_run(p4, ', now five restored historic properties on one walkable main street, with a spa on the hill above the lake, is one of the finest places to stay in New York State.')

# ============================================================
# AT A GLANCE TABLE
# ============================================================
add_heading_styled('At a glance', level=2)

table_data = [
    ('Length', '38.9 miles'),
    ('Depth', '435 feet'),
    ('Counties', 'Cayuga, Seneca, Tompkins'),
    ('Wine trail', '21 producers'),
    ('From NYC', '4 hours'),
    ('From Buffalo', '2 hours'),
    ('From Syracuse', '1 hour'),
    ('Anchor cities', 'Ithaca \u00b7 Auburn'),
]
table = doc.add_table(rows=len(table_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (label, value) in enumerate(table_data):
    row = table.rows[i]
    cell0 = row.cells[0]
    cell1 = row.cells[1]
    p0 = cell0.paragraphs[0]
    run0 = p0.add_run(label)
    run0.bold = True
    run0.font.name = 'Arial'
    run0.font.size = Pt(10)
    p1r = cell1.paragraphs[0]
    run1 = p1r.add_run(value)
    run1.font.name = 'Arial'
    run1.font.size = Pt(10)

print("Part 1 done: header through at-a-glance")

# ============================================================
# THE TOWNS
# ============================================================
add_heading_styled('The towns', level=2)

towns_intro = doc.add_paragraph()
add_run(towns_intro, 'Cayuga Lake is long enough to have meaningfully different communities at different points along its shore. ')
add_hyperlink(towns_intro, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(towns_intro, ' anchors the south with gorge trails and universities. ')
add_hyperlink(towns_intro, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(towns_intro, ' anchors the north with two national parks and the most consequential square mile in 19th-century American civil rights history. ')
add_hyperlink(towns_intro, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(towns_intro, ' holds the east shore\'s middle, modest in population and remarkable in biography.')

add_image_placeholder('IMAGE: Harriet Tubman National Historical Park \u2014 the home on South Street, Auburn NY')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Harriet Tubman', 'https://en.wikipedia.org/wiki/Harriet_Tubman', WIKI_BLUE)
add_run(cap, '\'s home at 180 South Street, ', italic=True)
add_hyperlink(cap, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(cap, ' \u2014 now the ', italic=True)
add_hyperlink(cap, 'Harriet Tubman National Historical Park', 'https://www.nps.gov/hart/', OTHER_PURPLE)
add_run(cap, '. She lived here from 1859 until her death in 1913.', italic=True)

# --- Auburn ---
h3 = add_heading_styled('Auburn', level=3)
meta = doc.add_paragraph()
add_run(meta, 'CAYUGA COUNTY \u00b7 NORTH END OF THE LAKE \u00b7 POP. ~26,000', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR THE REFORM CORRIDOR AND HARRIET TUBMAN', bold=True, size=10)

p = doc.add_paragraph()
add_hyperlink(p, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p, ' calls itself History\'s Hometown, which is the kind of claim that usually invites skepticism. Here it holds up. Within six blocks of downtown ')
add_hyperlink(p, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p, ', you have the home and burial site of ')
add_hyperlink(p, 'Harriet Tubman', 'https://en.wikipedia.org/wiki/Harriet_Tubman', WIKI_BLUE)
add_run(p, ', the home of ')
add_hyperlink(p, 'William H. Seward', 'https://en.wikipedia.org/wiki/William_H._Seward', WIKI_BLUE)
add_run(p, ', the only complete unaltered Tiffany interior in the United States, and the ')
add_hyperlink(p, 'New York State Equal Rights Heritage Center', 'https://www.google.com/maps/search/?api=1&query=New+York+State+Equal+Rights+Heritage+Center+Auburn+NY', MAPS_GREEN)
add_run(p, '. ')
add_hyperlink(p, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p, ' is where the reform movements of the 19th century were not visiting \u2014 they were living.')

p2 = doc.add_paragraph()
add_run(p2, 'The city sits at the northern end of ')
add_hyperlink(p2, 'Owasco Lake', 'https://www.google.com/maps/search/?api=1&query=Owasco+Lake+NY', MAPS_GREEN)
add_run(p2, ', one of the eastern ')
add_hyperlink(p2, 'Finger Lakes', 'https://en.wikipedia.org/wiki/Finger_Lakes', WIKI_BLUE)
add_run(p2, ', and is the economic center for the northern end of Cayuga. It has a working downtown, independent restaurants and breweries, and the weight of history distributed across it in a way that rewards walking rather than driving. March 10 is observed annually as ')
add_hyperlink(p2, 'Harriet Tubman', 'https://en.wikipedia.org/wiki/Harriet_Tubman', WIKI_BLUE)
add_run(p2, ' Day.')

# --- Aurora ---
add_heading_styled('Aurora', level=3)
meta = doc.add_paragraph()
add_run(meta, 'CAYUGA COUNTY \u00b7 EAST SHORE, MID-LAKE \u00b7 POP. ~700', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR THE EAST-SHORE OVERNIGHT AND THE VILLAGE WALK', bold=True, size=10)

p = doc.add_paragraph()
add_run(p, 'See the anchor identity section above \u2014 the full biography is there. What\'s worth adding for the town portrait: ')
add_hyperlink(p, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(p, ' is walkable in twenty minutes, and those twenty minutes pass the ')
add_hyperlink(p, 'Aurora Inn', 'https://www.innsofaurora.com/', OTHER_PURPLE)
add_run(p, ', the Fargo Bar, the ')
add_hyperlink(p, 'MacKenzie-Childs', 'https://www.mackenzie-childs.com/', OTHER_PURPLE)
add_run(p, ' village store, and a clear view of the lake from the elm-lined main street. ')
add_hyperlink(p, 'Long Point State Park', 'https://www.google.com/maps/search/?api=1&query=Long+Point+State+Park+Aurora+NY', MAPS_GREEN)
add_run(p, ' sits at the lake\'s edge just south of the village \u2014 free public water access, dock, picnic area.')

add_image_placeholder('IMAGE: Inns of Aurora \u2014 Aurora Inn facade or EB Morgan House with Cayuga Lake, east shore')
cap = doc.add_paragraph()
add_run(cap, 'The ', italic=True)
add_hyperlink(cap, 'Aurora Inn', 'https://www.innsofaurora.com/', OTHER_PURPLE)
add_run(cap, ', built 1833 by New York Times co-founder E.B. Morgan. One of five historic properties comprising the ', italic=True)
add_hyperlink(cap, 'Inns of Aurora', 'https://www.innsofaurora.com/', OTHER_PURPLE)
add_run(cap, '.', italic=True)

# --- Ithaca ---
add_heading_styled('Ithaca', level=3)
meta = doc.add_paragraph()
add_run(meta, 'TOMPKINS COUNTY \u00b7 SOUTH END OF THE LAKE \u00b7 POP. ~35,000 + 20,000 STUDENTS', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR GORGES, FOOD, AND THE SOUTHERN BASE', bold=True, size=10)

p = doc.add_paragraph()
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, ' sits at Cayuga\'s southern tip where the land rises sharply on both sides and the gorges begin. ')
add_hyperlink(p, 'Cornell University', 'https://en.wikipedia.org/wiki/Cornell_University', WIKI_BLUE)
add_run(p, '\'s alma mater \u2014 "Far Above Cayuga\'s Waters" \u2014 is considered the most widely copied university song in American higher education, which is a useful index of how far Cayuga Lake\'s reputation has traveled. ')
add_hyperlink(p, 'Ithaca College', 'https://en.wikipedia.org/wiki/Ithaca_College', WIKI_BLUE)
add_run(p, ' sits on the south hill. The two institutions give the city a food scene, arts community, and restaurant density \u2014 more restaurants per capita than New York City \u2014 that belies its size.')

p2 = doc.add_paragraph()
add_run(p2, 'The number that lands: more than 150 waterfalls within ten miles of downtown. The ')
add_hyperlink(p2, 'Cascadilla Gorge', 'https://www.google.com/maps/search/?api=1&query=Cascadilla+Gorge+Ithaca+NY', MAPS_GREEN)
add_run(p2, ' trail runs from the Commons to the ')
add_hyperlink(p2, 'Cornell', 'https://en.wikipedia.org/wiki/Cornell_University', WIKI_BLUE)
add_run(p2, ' campus \u2014 400 feet of elevation, six waterfalls, stone staircases built in the 1920s. It is also, for some ')
add_hyperlink(p2, 'Cornell', 'https://en.wikipedia.org/wiki/Cornell_University', WIKI_BLUE)
add_run(p2, ' students, just the commute. ')
add_hyperlink(p2, 'Stewart Park', 'https://www.google.com/maps/search/?api=1&query=Stewart+Park+Ithaca+NY', MAPS_GREEN)
add_run(p2, ' at the lake\'s edge south of the marina was formerly the site of the Wharton Brothers silent film studios \u2014 one of the earliest American film production centers before Hollywood. ')
add_hyperlink(p2, 'Rod Serling', 'https://en.wikipedia.org/wiki/Rod_Serling', WIKI_BLUE)
add_run(p2, ', creator of The Twilight Zone, summered in ')
add_hyperlink(p2, 'Interlaken', 'https://www.google.com/maps/search/?api=1&query=Interlaken+NY', MAPS_GREEN)
add_run(p2, ' on Cayuga\'s west shore.')

add_image_placeholder('IMAGE: Cascadilla Gorge \u2014 layered Devonian shale walls, trail, waterfall, Ithaca NY')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Cascadilla Gorge', 'https://www.google.com/maps/search/?api=1&query=Cascadilla+Gorge+Ithaca+NY', MAPS_GREEN)
add_run(cap, ', ', italic=True)
add_hyperlink(cap, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(cap, '. The trail drops 400 feet through Devonian shale from Cornell\'s campus to downtown. For some students, it\'s the commute.', italic=True)

# --- Seneca Falls ---
add_heading_styled('Seneca Falls', level=3)
meta = doc.add_paragraph()
add_run(meta, 'SENECA COUNTY \u00b7 NORTHWEST CORNER \u00b7 POP. ~6,800', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR THE WOMEN\'S RIGHTS STORY AND THE DECLARATION OF SENTIMENTS', bold=True, size=10)

p = doc.add_paragraph()
add_run(p, 'On July 19-20, 1848, 300 people gathered in the ')
add_hyperlink(p, 'Wesleyan Chapel', 'https://www.google.com/maps/search/?api=1&query=Wesleyan+Chapel+Seneca+Falls+NY', MAPS_GREEN)
add_run(p, ' here for the first formal women\'s rights convention in American history. ')
add_hyperlink(p, 'Elizabeth Cady Stanton', 'https://en.wikipedia.org/wiki/Elizabeth_Cady_Stanton', WIKI_BLUE)
add_run(p, ' read the Declaration of Sentiments, modeled phrase for phrase on Jefferson\'s Declaration of Independence. ')
add_hyperlink(p, 'Frederick Douglass', 'https://en.wikipedia.org/wiki/Frederick_Douglass', WIKI_BLUE)
add_run(p, ' \u2014 the only Black person in attendance \u2014 made the argument that clinched the vote on women\'s suffrage. The resolution passed. It took 72 more years.')

p2 = doc.add_paragraph()
add_run(p2, 'The ')
add_hyperlink(p2, 'Women\'s Rights National Historical Park', 'https://en.wikipedia.org/wiki/Women%27s_Rights_National_Historical_Park', WIKI_BLUE)
add_run(p2, ' encompasses the chapel, the M\'Clintock House in ')
add_hyperlink(p2, 'Waterloo', 'https://www.google.com/maps/search/?api=1&query=Waterloo+NY', MAPS_GREEN)
add_run(p2, ' where the Declaration was drafted, and Stanton\'s home. The town also claims, with varying degrees of scholarly support, to be the inspiration for Bedford Falls in It\'s a Wonderful Life. The claim has been disputed for decades. ')
add_hyperlink(p2, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(p2, ' is clearly winning on tourism grounds, which is itself a very Seneca Falls thing to have pulled off.')

add_image_placeholder('IMAGE: Wesleyan Chapel Seneca Falls \u2014 Women\'s Rights National Historical Park exterior')
cap = doc.add_paragraph()
add_run(cap, 'The ', italic=True)
add_hyperlink(cap, 'Wesleyan Chapel', 'https://www.google.com/maps/search/?api=1&query=Wesleyan+Chapel+Seneca+Falls+NY', MAPS_GREEN)
add_run(cap, ', ', italic=True)
add_hyperlink(cap, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(cap, '. Site of the 1848 Women\'s Rights Convention. Now part of the ', italic=True)
add_hyperlink(cap, 'Women\'s Rights National Historical Park', 'https://www.nps.gov/wori/', OTHER_PURPLE)
add_run(cap, '.', italic=True)

# --- Trumansburg ---
add_heading_styled('Trumansburg', level=3)
meta = doc.add_paragraph()
add_run(meta, 'TOMPKINS COUNTY \u00b7 WEST SHORE, MID-LAKE \u00b7 POP. ~1,500', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR THE WEST-SHORE DAY AND TAUGHANNOCK FALLS', bold=True, size=10)

p = doc.add_paragraph()
add_run(p, 'Two miles south of ')
add_hyperlink(p, 'Taughannock Falls State Park', 'https://en.wikipedia.org/wiki/Taughannock_Falls_State_Park', WIKI_BLUE)
add_run(p, ' and two miles north of the ')
add_hyperlink(p, 'Finger Lakes Cider House', 'https://www.fingerlakesciderhouse.com/', OTHER_PURPLE)
add_run(p, ' at ')
add_hyperlink(p, 'Good Life Farm', 'https://www.google.com/maps/search/?api=1&query=Good+Life+Farm+Interlaken+NY', MAPS_GREEN)
add_run(p, '. That geography makes it the anchor for a west-shore day. Small main street, Atlas Bowl for dinner \u2014 elevated comfort food without performing it. The kind of town that rewards people who stop.')

# --- Union Springs ---
add_heading_styled('Union Springs', level=3)
meta = doc.add_paragraph()
add_run(meta, 'CAYUGA COUNTY \u00b7 EAST SHORE, UPPER MID-LAKE \u00b7 POP. ~1,100', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR SMALL-PRODUCTION RIESLING AND LAKEFRONT DINING', bold=True, size=10)

p = doc.add_paragraph()
add_hyperlink(p, 'Heart & Hands Wine Company', 'https://www.heartandhandswine.com/', OTHER_PURPLE)
add_run(p, ' is here \u2014 one of the trail\'s most focused producers, small-production Riesling from the east shore. The ')
add_hyperlink(p, 'Wheelhouse Restaurant', 'https://www.google.com/maps/search/?api=1&query=Wheelhouse+Restaurant+Union+Springs+NY', MAPS_GREEN)
add_run(p, ', third-generation family-owned, has been on the lake since the late 1970s.')

print("Part 2 done: towns section")

# ============================================================
# THE HERITAGE CORRIDOR (renamed from "The reform corridor")
# ============================================================
add_heading_styled('The Heritage Corridor', level=2)

add_image_placeholder('IMAGE: Harriet Tubman portrait \u2014 LOC public domain 1868 photograph')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Harriet Tubman', 'https://en.wikipedia.org/wiki/Harriet_Tubman', WIKI_BLUE)
add_run(cap, ', c. 1868. She moved to ', italic=True)
add_hyperlink(cap, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(cap, ' in 1859 and spent 54 years organizing there. Library of Congress.', italic=True)

p = doc.add_paragraph()
add_run(p, 'The fifteen miles between ')
add_hyperlink(p, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p, ' and ')
add_hyperlink(p, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(p, ' contain two national parks, the Underground Railroad\'s most consequential conductor, and the founding document of the American women\'s rights movement. They also contain the only complete unaltered Tiffany interior in the United States, which is the kind of detail that only gets mentioned in a guide that\'s actually paying attention.')

p2 = doc.add_paragraph()
add_run(p2, 'The reform movements that produced all of this were not separate campaigns that happened to overlap geographically. They were the same network \u2014 the same Quaker meeting houses, the same abolitionist families, the same roads. ')
add_hyperlink(p2, 'Harriet Tubman', 'https://en.wikipedia.org/wiki/Harriet_Tubman', WIKI_BLUE)
add_run(p2, ' and ')
add_hyperlink(p2, 'Elizabeth Cady Stanton', 'https://en.wikipedia.org/wiki/Elizabeth_Cady_Stanton', WIKI_BLUE)
add_run(p2, ' knew each other. ')
add_hyperlink(p2, 'Frederick Douglass', 'https://en.wikipedia.org/wiki/Frederick_Douglass', WIKI_BLUE)
add_run(p2, ' was a guest at the ')
add_hyperlink(p2, 'Seward House', 'https://www.sewardhouse.org/', OTHER_PURPLE)
add_run(p2, ' in ')
add_hyperlink(p2, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p2, ' and a speaker at the ')
add_hyperlink(p2, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(p2, ' convention. The geography held the movement together.')

p3 = doc.add_paragraph()
add_run(p3, 'To this day, the ')
add_hyperlink(p3, 'Haudenosaunee', 'https://en.wikipedia.org/wiki/Iroquois', WIKI_BLUE)
add_run(p3, ' refer to the office of the President of the United States as Hanadaga\u0301:yas \u2014 \'He Who Destroys Villages.\' The name dates to 1779. The Cayuga people are still here.')

# --- Haudenosaunee subsection ---
add_heading_styled('The Haudenosaunee \u2014 where it begins', level=3)

p = doc.add_paragraph()
add_run(p, 'The ')
add_hyperlink(p, 'Cayuga Nation', 'https://en.wikipedia.org/wiki/Cayuga_Nation', WIKI_BLUE)
add_run(p, ', Gayogo\u0331ho\u0301:n\u01eb\u02bc in their own language, are the People of the Great Swamp \u2014 one of the founding nations of the ')
add_hyperlink(p, 'Haudenosaunee Confederacy', 'https://en.wikipedia.org/wiki/Iroquois', WIKI_BLUE)
add_run(p, '. They held the territory around Cayuga Lake for centuries, documented at ')
add_hyperlink(p, 'Union Springs', 'https://www.google.com/maps/search/?api=1&query=Union+Springs+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(p, ', Cayuga village, ')
add_hyperlink(p, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(p, ', and ')
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, ', sustained by the same agricultural conditions that make the land productive today.')

p2 = doc.add_paragraph()
add_run(p2, 'The ')
add_hyperlink(p2, 'Sullivan-Clinton Campaign', 'https://en.wikipedia.org/wiki/Sullivan%E2%80%93Clinton_campaign', WIKI_BLUE)
add_run(p2, ' burned those towns in 1779. Washington\'s orders were explicit: total destruction, ruin the crops. The Cayuga were dispossessed under treaties negotiated in the 1790s that the Nation regards as illegitimate to this day. The lake still has their name. In 2005, a 70-acre farm in ')
add_hyperlink(p2, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(p2, ' \u2014 on the land of Chonodote itself \u2014 was signed over to the ')
add_hyperlink(p2, 'Cayuga Nation', 'https://en.wikipedia.org/wiki/Cayuga_Nation', WIKI_BLUE)
add_run(p2, ' by US citizens who had purchased it.')

add_image_placeholder('IMAGE: Sullivan-Clinton Campaign 1779 \u2014 NPS or sullivanclinton.com historical map')
cap = doc.add_paragraph()
add_run(cap, 'Map of the ', italic=True)
add_hyperlink(cap, 'Sullivan-Clinton Campaign', 'https://en.wikipedia.org/wiki/Sullivan%E2%80%93Clinton_campaign', WIKI_BLUE)
add_run(cap, ', 1779. Continental Army forces moved along Cayuga\'s western shore, burning the Cayuga towns including Chonodote at present-day ', italic=True)
add_hyperlink(cap, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(cap, '. NPS / public domain.', italic=True)

# --- Harriet Tubman in Auburn ---
add_heading_styled('Harriet Tubman in Auburn', level=3)

p = doc.add_paragraph()
add_hyperlink(p, 'Tubman', 'https://en.wikipedia.org/wiki/Harriet_Tubman', WIKI_BLUE)
add_run(p, ' chose ')
add_hyperlink(p, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p, ' in 1859 because of what was already there: a Quaker abolitionist network, a sympathetic senator two blocks away, and proximity to the Canadian border routes she knew from conducting. She moved her aging parents from Canada and settled on a seven-acre farm at 180 South Street in ')
add_hyperlink(p, 'Fleming', 'https://www.google.com/maps/search/?api=1&query=Fleming+NY', MAPS_GREEN)
add_run(p, ', just outside the city line. The farm was sold to her by ')
add_hyperlink(p, 'William H. Seward', 'https://en.wikipedia.org/wiki/William_H._Seward', WIKI_BLUE)
add_run(p, ' \u2014 then a U.S. Senator \u2014 who held the mortgage privately and was flexible about payments. Tubman was legally a fugitive. She had no property rights under U.S. law. Seward did it anyway.')

p2 = doc.add_paragraph()
add_run(p2, 'She spent 54 years in ')
add_hyperlink(p2, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p2, ' after that \u2014 organizing, supporting civil rights and women\'s suffrage, caring for people who had nowhere else to go. In 1908 she built the Home for the Aged at 182 South Street, which she deeded to the Thompson AME Zion Church. She died there in 1913 at approximately 91. The grave at ')
add_hyperlink(p2, 'Fort Hill Cemetery', 'https://www.google.com/maps/search/?api=1&query=Fort+Hill+Cemetery+Auburn+NY', MAPS_GREEN)
add_run(p2, ' is three blocks from her home. The marker is famously spare: just her name.')

p3 = doc.add_paragraph()
add_run(p3, 'The ')
add_hyperlink(p3, 'Harriet Tubman National Historical Park', 'https://www.nps.gov/hart/', OTHER_PURPLE)
add_run(p3, ', established in 2017, encompasses the residence, the Home for the Aged, and the Thompson AME Zion Church. The NPS visitor center at the church is the best starting point.')

# --- The Seward House ---
add_heading_styled('The Seward House', level=3)

p = doc.add_paragraph()
add_hyperlink(p, 'William Henry Seward', 'https://en.wikipedia.org/wiki/William_H._Seward', WIKI_BLUE)
add_run(p, ' served as New York Governor, U.S. Senator, and Secretary of State under both Lincoln and Johnson \u2014 the man who negotiated the purchase of Alaska, dismissed at the time as Seward\'s Folly. His house on South Street in ')
add_hyperlink(p, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p, ' was a documented Underground Railroad stop, a meeting point for the reform networks, and the place where Seward privately arranged the land transaction with ')
add_hyperlink(p, 'Tubman', 'https://en.wikipedia.org/wiki/Harriet_Tubman', WIKI_BLUE)
add_run(p, '. ')
add_hyperlink(p, 'Frederick Douglass', 'https://en.wikipedia.org/wiki/Frederick_Douglass', WIKI_BLUE)
add_run(p, ' was a frequent visitor. The ')
add_hyperlink(p, 'Seward House Museum', 'https://www.sewardhouse.org/', OTHER_PURPLE)
add_run(p, ', operated since 1955, preserves the original furnishings and the physical weight of the 19th century still present in its rooms.')

# --- Willard Memorial Chapel ---
add_heading_styled('The Willard Memorial Chapel', level=3)

p = doc.add_paragraph()
add_run(p, 'The ')
add_hyperlink(p, 'Willard Memorial Chapel', 'https://www.google.com/maps/search/?api=1&query=Willard+Memorial+Chapel+Auburn+NY', MAPS_GREEN)
add_run(p, ' on Nelson Street in ')
add_hyperlink(p, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p, ' is the only complete, unaltered ')
add_hyperlink(p, 'Louis Comfort Tiffany', 'https://en.wikipedia.org/wiki/Louis_Comfort_Tiffany', WIKI_BLUE)
add_run(p, ' interior remaining in the United States. Tiffany did the entire space in 1894 \u2014 the glass, the tiles, the metalwork, the furnishings. Everything. It survived because it was attached to an institution that closed and transferred ownership carefully rather than demolishing. Most people driving through ')
add_hyperlink(p, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(p, ' don\'t know it exists.')

add_image_placeholder('IMAGE: Willard Memorial Chapel interior \u2014 Tiffany glass and decorating, Auburn NY 1894')
cap = doc.add_paragraph()
add_run(cap, 'The ', italic=True)
add_hyperlink(cap, 'Willard Memorial Chapel', 'https://www.google.com/maps/search/?api=1&query=Willard+Memorial+Chapel+Auburn+NY', MAPS_GREEN)
add_run(cap, ', ', italic=True)
add_hyperlink(cap, 'Auburn', 'https://www.google.com/maps/search/?api=1&query=Auburn+NY', MAPS_GREEN)
add_run(cap, '. The only complete, unaltered ', italic=True)
add_hyperlink(cap, 'Louis Comfort Tiffany', 'https://en.wikipedia.org/wiki/Louis_Comfort_Tiffany', WIKI_BLUE)
add_run(cap, ' interior in the United States. 1894.', italic=True)

add_image_placeholder('IMAGE: Elizabeth Cady Stanton portrait \u2014 1870 Napoleon Sarony photograph, public domain')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Elizabeth Cady Stanton', 'https://en.wikipedia.org/wiki/Elizabeth_Cady_Stanton', WIKI_BLUE)
add_run(cap, ', c. 1870. Primary author of the Declaration of Sentiments. She lived in ', italic=True)
add_hyperlink(cap, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(cap, ' and organized the 1848 convention with Lucretia Mott. Library of Congress.', italic=True)

# --- Heritage sites full list ---
add_heading_styled('Heritage sites \u2014 full list', level=3)

badge_p = doc.add_paragraph()
add_run(badge_p, 'Badge track: All sites below are tagged to the ', italic=True)
add_hyperlink(badge_p, 'Finger Lakes Heritage Trail', 'https://upstate.tourismo.app/trails/finger-lakes-heritage-trail', UPSTATE_RED)
add_run(badge_p, ' passport badge.', italic=True)

heritage_sites = [
    ('Harriet Tubman National Historical Park', 'Auburn / Fleming', 'Where freedom found an address', 'The home, the Home for the Aged, the Thompson AME Zion Church. NPS visitor center at the church. Tubman lived at 180 South Street from 1859 until her death in 1913, making it the longest-occupied site in the Underground Railroad network. The Home for the Aged, which she built in 1908 and deeded to the Thompson AME Zion Church, served elderly African Americans with nowhere else to go. The NPS visitor center at the church is the best starting point; rangers lead guided walks to the residence and the Home for the Aged. Open year-round. Free admission.', 'https://www.nps.gov/hart/', 'https://www.google.com/maps/search/?api=1&query=Harriet+Tubman+National+Historical+Park+Auburn+NY'),
    ('Fort Hill Cemetery', 'Auburn', 'The marker is famously spare: just her name', 'Tubman\'s burial site, three blocks from her home. Designated as a public park in 1836, the cemetery also holds the graves of William H. Seward and other figures from Auburn\'s reform era. The Tubman gravesite is marked with a simple headstone. Path Through History site.', None, 'https://www.google.com/maps/search/?api=1&query=Fort+Hill+Cemetery+Auburn+NY'),
    ('Seward House Museum', 'Auburn', 'The senator who sold Tubman a farm and bought Alaska', 'Underground Railroad documented stop and the private residence of William Henry Seward, who served as Governor, Senator, and Secretary of State. Original furnishings, documents, and the room where Seward arranged the land sale to Tubman. Frederick Douglass was a frequent guest. Operated as a museum since 1955. Path Through History site.', 'https://www.sewardhouse.org/', 'https://www.google.com/maps/search/?api=1&query=Seward+House+Museum+Auburn+NY'),
    ('NYS Equal Rights Heritage Center', 'Auburn', 'The reform network, mapped and explained', 'Interactive exhibits on Tubman, Stanton, Douglass, and the overlapping reform movements that ran through this corridor. The center connects the heritage sites into a single narrative and serves as a practical starting point for the Auburn walking itinerary. Free admission.', None, 'https://www.google.com/maps/search/?api=1&query=New+York+State+Equal+Rights+Heritage+Center+Auburn+NY'),
    ('Willard Memorial Chapel', 'Auburn', 'The last complete Tiffany interior on earth', 'The only complete, unaltered Louis Comfort Tiffany interior remaining in the United States. Tiffany designed the entire space in 1894: the stained glass, the mosaic floor tiles, the metalwork, the oak furnishings. It survived because the adjoining seminary closed and transferred ownership carefully rather than demolishing. Most people driving through Auburn don\'t know it exists. Path Through History site.', None, 'https://www.google.com/maps/search/?api=1&query=Willard+Memorial+Chapel+Auburn+NY'),
    ('Women\'s Rights National Historical Park', 'Seneca Falls', 'Where 300 people rewrote the rules in two days', 'The Wesleyan Chapel where 300 people gathered on July 19-20, 1848 for the first formal women\'s rights convention in American history. Elizabeth Cady Stanton read the Declaration of Sentiments; Frederick Douglass made the argument that clinched the suffrage vote. Also encompasses the M\'Clintock House in Waterloo where the Declaration was drafted and Stanton\'s home. NPS visitor center on Fall Street.', 'https://www.nps.gov/wori/', 'https://www.google.com/maps/search/?api=1&query=Womens+Rights+National+Historical+Park+Seneca+Falls+NY'),
    ('Howland Stone Store Museum', 'Aurora', 'Quaker abolitionists, cobblestone walls, 1837', '1837 cobblestone commercial building on Aurora\'s main street. The Howland family were prominent Quaker abolitionists connected to the Underground Railroad network. The museum preserves the building and tells the story of the reform-era families who made Aurora a node in the abolitionist movement decades before Tubman arrived.', None, 'https://www.google.com/maps/search/?api=1&query=Howland+Stone+Store+Museum+Aurora+NY'),
    ('Montezuma Heritage Park', 'Montezuma', 'The crossing point at the lake\'s north end', 'Along the Seneca River and Erie Canal at Cayuga Lake\'s northern outlet. The park marks a documented Underground Railroad crossing where the canal and river systems provided routes north toward Canada. The landscape connects the heritage corridor to the waterways that made the network function.', None, 'https://www.google.com/maps/search/?api=1&query=Montezuma+Heritage+Park+NY'),
    ('Museum of the Earth', 'Ithaca', 'Two million years of glacial history told in stone', 'Paleontological research museum affiliated with Cornell. The permanent collection explains the glacial formation that carved the Finger Lakes, created the gorges, and deposited the shale soils that make the wine region possible. The Right Whale skeleton in the main hall is 40 feet long. A useful first stop for visitors who want to understand why the landscape looks the way it does.', 'https://www.museumoftheearth.org/', 'https://www.google.com/maps/search/?api=1&query=Museum+of+the+Earth+Ithaca+NY'),
    ('Ward O\'Hara Agricultural Museum', 'Auburn', 'Known for connecting farming past to the contemporary trail', 'Regional agricultural history spanning two centuries of Finger Lakes farming. The collection connects the practices of the Haudenosaunee and early settlers to the contemporary farm trail. Useful context for the Cayuga Farm Loop.', None, 'https://www.google.com/maps/search/?api=1&query=Ward+O+Hara+Agricultural+Museum+Auburn+NY'),
]

for name, location, subtitle, desc, official_url, maps_url in heritage_sites:
    p = doc.add_paragraph()
    add_run(p, '\u2014 ')
    if official_url:
        add_hyperlink(p, name, official_url, OTHER_PURPLE)
    else:
        add_run(p, name, bold=True)
    add_run(p, ' ')
    add_hyperlink(p, location, maps_url, MAPS_GREEN)
    sub_p = doc.add_paragraph()
    add_run(sub_p, subtitle, italic=True)
    doc.add_paragraph(desc)

print("Part 3 done: heritage corridor")

# ============================================================
# THE CAYUGA LAKE WINE TRAIL
# ============================================================
add_heading_styled('The Cayuga Lake Wine Trail', level=2)

p = doc.add_paragraph()
add_run(p, 'The ')
add_hyperlink(p, 'Cayuga Lake Wine Trail', 'https://upstate.tourismo.app/trails/cayuga-lake-wine-trail', UPSTATE_RED)
add_run(p, ' was the first organized wine trail in the United States, established in 1983. There are 21 producers in the current passport, spanning all three counties that touch the lake. The trail runs Route 89 on the west shore \u2014 the primary loop \u2014 and Route 90 on the east shore, which is quieter and produces some of the most interesting bottles on the trail.')

p2 = doc.add_paragraph()
add_run(p2, 'What distinguishes Cayuga from Seneca: higher elevation, steeper shale slopes, and a slightly more austere expression in the Riesling \u2014 more linear and mineral, more transparently terroir-driven. The east shore in particular produces Rieslings that invite serious comparison with the wines they were modeled after in the ')
add_hyperlink(p2, 'Mosel', 'https://en.wikipedia.org/wiki/Moselle', WIKI_BLUE)
add_run(p2, '.')

add_image_placeholder('IMAGE: Cayuga Lake Wine Trail \u2014 vineyard hillside, lake visible below, east shore preferred')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Cayuga Lake Wine Trail', 'https://upstate.tourismo.app/trails/cayuga-lake-wine-trail', UPSTATE_RED)
add_run(cap, '. The east shore\'s shale slopes produce some of the most precisely cold-climate Rieslings in New York State.', italic=True)

# Producers
add_heading_styled('Producers worth naming', level=3)

producers = [
    ('Heart & Hands Wine Company', 'Union Springs \u2014 east shore', 'The east shore at its most focused', 'Small-production Riesling specialist. Tom and Susan Higgins farm a single east-shore vineyard and make wine with a restraint that lets the site speak. The tasting room is intimate and unhurried. If you want to understand what Cayuga east-shore Riesling tastes like at its most precise, this is the stop.', 'https://www.heartandhandswine.com/', 'https://www.google.com/maps/search/?api=1&query=Heart+and+Hands+Wine+Company+Union+Springs+NY'),
    ('Treleaven Wines', 'King Ferry \u2014 east shore', 'Thirty years on the same hillside', 'Thirty-year estate operation on the east shore. Award-winning dry Riesling from vines that have had decades to find their depth. Outdoor terrace overlooking the lake. The kind of producer that doesn\'t need to explain itself because the wine does the talking.', 'https://www.treleavenwines.com/', 'https://www.google.com/maps/search/?api=1&query=Treleaven+Wines+King+Ferry+NY'),
    ('Boundary Breaks Vineyard', 'Lodi \u2014 east Cayuga / Seneca border', 'One vineyard, one variety, no compromises', 'Single-vineyard Riesling program on the ridge between Cayuga and Seneca Lakes. Wine Enthusiast named their Dry Riesling to its Top 100 Wines in the World twice. No buses, no groups larger than six. The argument for why single-vineyard designations matter in this AVA.', 'https://www.boundarybreaks.com/', 'https://www.google.com/maps/search/?api=1&query=Boundary+Breaks+Vineyard+Lodi+NY'),
    ('Americana Vineyards', 'Interlaken \u2014 west shore', 'Known for longevity and a family welcome', 'One of the older Cayuga producers, family-run since the beginning. Wine and cider on the same property, a broad portfolio that serves first-time trail visitors and repeat customers equally well. The west-shore stop with the widest range.', 'https://www.americanavineyards.com/', 'https://www.google.com/maps/search/?api=1&query=Americana+Vineyards+Interlaken+NY'),
    ('Buttonwood Grove Winery', 'Romulus \u2014 west shore', 'Estate-grown wine, lakeside pour', 'Estate-grown grapes, lakeside tasting room on the west shore. The vineyard faces the lake directly, and the tasting room takes advantage of it. A straightforward, well-run operation that rewards a stop on the Route 89 loop.', 'https://www.buttonwoodgrove.com/', 'https://www.google.com/maps/search/?api=1&query=Buttonwood+Grove+Winery+Romulus+NY'),
    ('Frontenac Point Vineyard', 'Trumansburg \u2014 west shore', 'The quiet stop near Taughannock', 'Family estate on the west shore near Taughannock Falls State Park. Small production, personal pours, the kind of stop where the winemaker might be the one pouring. One of the trail\'s quieter experiences, best paired with a Taughannock visit.', 'https://www.frontenacpoint.com/', 'https://www.google.com/maps/search/?api=1&query=Frontenac+Point+Vineyard+Trumansburg+NY'),
    ('Six Mile Creek Vineyards', 'Ithaca', 'The closest winery to the Commons', 'Closest winery to downtown Ithaca, making it the natural entry point for first-time visitors and the last stop before heading back to town. Broad portfolio, casual tasting room, outdoor seating with valley views.', 'https://www.sixmilecreek.com/', 'https://www.google.com/maps/search/?api=1&query=Six+Mile+Creek+Vineyards+Ithaca+NY'),
    ('Long Point Winery', 'Aurora \u2014 east shore', 'Best paired with a night at the Inns', 'East-shore tasting room in Aurora, worth combining with a stay at the Inns of Aurora for the complete east-shore day. The view from the tasting room faces west across the lake. Small-scale, personal service.', 'https://www.longpointwinery.com/', 'https://www.google.com/maps/search/?api=1&query=Long+Point+Winery+Aurora+NY'),
]

for name, location, subtitle, desc, official_url, maps_url in producers:
    p = doc.add_paragraph()
    add_run(p, '\u2014 ')
    add_hyperlink(p, name, official_url, OTHER_PURPLE)
    add_run(p, ' ')
    add_hyperlink(p, location, maps_url, MAPS_GREEN)
    sub_p = doc.add_paragraph()
    add_run(sub_p, subtitle, italic=True)
    doc.add_paragraph(desc)

# How to run the trail
add_heading_styled('How to run the trail', level=3)

p = doc.add_paragraph()
add_run(p, 'The standard west-shore loop runs Route 89 from ')
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, ' north to Cayuga village. A full circuit takes two days done properly \u2014 one shore per day, with a night in ')
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, ' or ')
add_hyperlink(p, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(p, ' between. The east shore is the better second-day drive: fewer stops, more space, producers who see fewer visitors and have more time to talk. Tasting rooms run full hours May through November; call ahead November through April.')

# ============================================================
# FARMS, CIDER, AND FOOD
# ============================================================
add_heading_styled('Farms, cider, and food', level=2)

# --- Finger Lakes Cider House ---
farm_p = doc.add_paragraph()
add_run(farm_p, '\u2014 ')
add_hyperlink(farm_p, 'Finger Lakes Cider House at Good Life Farm', 'https://www.fingerlakesciderhouse.com/', OTHER_PURPLE)
add_run(farm_p, ' ')
add_hyperlink(farm_p, 'Interlaken / Trumansburg', 'https://www.google.com/maps/search/?api=1&query=Good+Life+Farm+Interlaken+NY', MAPS_GREEN)
sub_p = doc.add_paragraph()
add_run(sub_p, 'Four cideries under one roof on a working organic farm', italic=True)
doc.add_paragraph('The anchor of the Cayuga Farm Loop. Four cideries share a single tasting room on a working organic farm on Cayuga\'s western shore outside Trumansburg. Food is sourced from neighboring properties. The farm itself is open to walk; the cideries rotate their offerings seasonally. Pair with a stop at Lively Run Dairy two miles south for the full farm corridor experience. Badge: Finger Lakes Farm Trail.')

# --- Lively Run Dairy ---
farm_p = doc.add_paragraph()
add_run(farm_p, '\u2014 ')
add_hyperlink(farm_p, 'Lively Run Dairy', 'https://www.livelyrun.com/', OTHER_PURPLE)
add_run(farm_p, ' ')
add_hyperlink(farm_p, 'Interlaken', 'https://www.google.com/maps/search/?api=1&query=Lively+Run+Dairy+Interlaken+NY', MAPS_GREEN)
sub_p = doc.add_paragraph()
add_run(sub_p, 'Goat cheese since 1982, before artisan was a category', italic=True)
doc.add_paragraph('Lively Run has been making goat cheese in Interlaken since 1982, which makes it one of the oldest artisan creameries in the Finger Lakes. The farm store sells the full line; the goats are visible from the parking area. Two miles south of the Finger Lakes Cider House on Route 89. Badge: Finger Lakes Farm Trail.')

# --- Ithaca Farmers Market ---
farm_p = doc.add_paragraph()
add_run(farm_p, '\u2014 ')
add_hyperlink(farm_p, 'Ithaca Farmers Market', 'https://www.ithacamarket.com/', OTHER_PURPLE)
add_run(farm_p, ' ')
add_hyperlink(farm_p, 'Steamboat Landing, Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+Farmers+Market+Steamboat+Landing', MAPS_GREEN)
sub_p = doc.add_paragraph()
add_run(sub_p, 'The single best food destination on the lake', italic=True)
doc.add_paragraph('Saturday and Sunday at Steamboat Landing, producers-only, within sight of Cayuga Lake. Over 130 active vendors, everything grown or made within 30 miles of the pavilion. The market runs April through December on Saturdays, May through November on Sundays. Come early on Saturday; the parking lot fills by 10am.')

# --- Route 89 Farm Stands ---
farm_p = doc.add_paragraph()
add_run(farm_p, '\u2014 Route 89 Farm Stands')
add_run(farm_p, ' ', bold=True)
add_hyperlink(farm_p, 'West shore, Ithaca to Romulus', 'https://www.google.com/maps/search/?api=1&query=Route+89+Cayuga+Lake+NY', MAPS_GREEN)
sub_p = doc.add_paragraph()
add_run(sub_p, 'Known for the farm corridor that predates the wine trail', italic=True)
doc.add_paragraph('The farm stands along Route 89 on Cayuga\'s west shore have been there longer than the wine trail designation. June through October, the road between Ithaca and Romulus is lined with seasonal produce, cut flowers, and honey. No single stand dominates; the cumulative effect is the point. Badge: Finger Lakes Farm Trail.')

add_image_placeholder('IMAGE: Finger Lakes Cider House at Good Life Farm \u2014 tasting room exterior, farm setting, Cayuga Lake in background')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Finger Lakes Cider House', 'https://www.fingerlakesciderhouse.com/', OTHER_PURPLE)
add_run(cap, ' at ', italic=True)
add_hyperlink(cap, 'Good Life Farm', 'https://www.google.com/maps/search/?api=1&query=Good+Life+Farm+Interlaken+NY', MAPS_GREEN)
add_run(cap, ', ', italic=True)
add_hyperlink(cap, 'Interlaken', 'https://www.google.com/maps/search/?api=1&query=Interlaken+NY', MAPS_GREEN)
add_run(cap, '. Four cideries under one roof on a working organic farm on Route 89.', italic=True)

add_image_placeholder('IMAGE: Ithaca Farmers Market \u2014 Steamboat Landing, Cayuga Lake visible, vendors and pavilion')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Ithaca Farmers Market', 'https://www.ithacamarket.com/', OTHER_PURPLE)
add_run(cap, ', ', italic=True)
add_hyperlink(cap, 'Steamboat Landing', 'https://www.google.com/maps/search/?api=1&query=Ithaca+Farmers+Market+Steamboat+Landing', MAPS_GREEN)
add_run(cap, '. Saturday and Sunday, producers-only, over 130 vendors within 30 miles.', italic=True)

# ============================================================
# WHERE TO EAT (preserved as-is, excluded from this pass)
# ============================================================
add_heading_styled('Where to eat', level=2)

# Auburn restaurants
eat_auburn = doc.add_paragraph()
add_run(eat_auburn, 'AUBURN', bold=True)

restaurants_auburn = [
    ('Moro\'s Table', 'Auburn', 'Chef Moro sources locally. The best independent restaurant in Auburn.'),
    ('Elderberry Pond Restaurant', 'Auburn', 'Farm-to-table on a working organic farm south of the city.'),
    ('Refinery', 'Auburn', 'Locally sourced modern fare, good cocktail program. Downtown Auburn.'),
    ('Osteria Salina', 'Auburn', 'Upscale Italian. Good wine program. The dinner reservation option in Auburn.'),
    ('Cafe 108', 'Auburn', 'Award-winning cafe supporting Auburn Public Theater. Community institution.'),
    ('Prison City Pub & Brewery', 'Auburn', 'Named for Auburn\'s history. Local beer, downtown.'),
]
for name, loc, desc in restaurants_auburn:
    p = doc.add_paragraph()
    add_run(p, '\u2014 ')
    add_run(p, name, bold=True)
    add_run(p, ' ' + desc)

eat_aurora = doc.add_paragraph()
add_run(eat_aurora, 'AURORA AND EAST SHORE', bold=True)

restaurants_aurora = [
    ('1833 Kitchen & Bar', 'Aurora', 'Farm-to-table at Inns of Aurora. Chef Eric Lamphere. Locally sourced, lake views.'),
    ('The Fargo Bar & Grill', 'Aurora', 'E.B. Morgan\'s 1834 office building, now the Inns of Aurora\'s pub. Burgers, local beer.'),
    ('Wheelhouse Restaurant', 'Union Springs', 'Third-generation family ownership. On the lake. Casual lakefront dining.'),
]
for name, loc, desc in restaurants_aurora:
    p = doc.add_paragraph()
    add_run(p, '\u2014 ')
    add_run(p, name, bold=True)
    add_run(p, ' ' + desc)

eat_ithaca = doc.add_paragraph()
add_run(eat_ithaca, 'ITHACA AND SURROUNDS', bold=True)

restaurants_ithaca = [
    ('Moosewood Restaurant', 'Ithaca', 'Open since 1973. The country\'s longest-running vegetarian restaurant. Daily-changing menus, everything local.'),
    ('Gola Osteria', 'Ithaca', 'Elevated Italian in the Quarry Arms building. Housemade pasta, strong wine list.'),
    ('Ithaca Beer Co.', 'Ithaca', 'Flower Power IPA. Wood-fired food, valley views, outdoor space.'),
    ('Gimme! Coffee', 'Ithaca', 'Ithaca institution. Locally roasted. The standard morning stop.'),
    ('Atlas Bowl', 'Trumansburg', 'Elevated comfort food near Taughannock Falls. The lunch anchor for the outdoor day.'),
    ('Finger Lakes Cider House', 'Interlaken', 'Tasting room food sourced from neighboring farms. Cheese boards, seasonal small plates.'),
]
for name, loc, desc in restaurants_ithaca:
    p = doc.add_paragraph()
    add_run(p, '\u2014 ')
    add_run(p, name, bold=True)
    add_run(p, ' ' + desc)

print("Part 4 done: wine trail, farms, restaurants")

# ============================================================
# OUTDOOR AND NATURE (renamed from "Outside")
# ============================================================
add_heading_styled('Outdoor and nature', level=2)

p = doc.add_paragraph()
add_run(p, 'More than 150 waterfalls within ten miles of downtown ')
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, '. The standard recommendation is ')
add_hyperlink(p, 'Taughannock Falls', 'https://en.wikipedia.org/wiki/Taughannock_Falls_State_Park', WIKI_BLUE)
add_run(p, ' because it\'s easy to understate: 215 feet of vertical drop, three stories taller than Niagara, accessible via a flat 0.9-mile walk from the parking area. Go in April when the snowmelt is running and the sound reaches the trailhead.')

add_image_placeholder('IMAGE: Taughannock Falls \u2014 full 215-foot drop visible, gorge walls, preferably spring or early summer')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Taughannock Falls', 'https://en.wikipedia.org/wiki/Taughannock_Falls_State_Park', WIKI_BLUE)
add_run(cap, ', ', italic=True)
add_hyperlink(cap, 'Trumansburg', 'https://www.google.com/maps/search/?api=1&query=Trumansburg+NY', MAPS_GREEN)
add_run(cap, '. 215 feet \u2014 three stories taller than Niagara. The gorge trail from the parking area is flat and under a mile.', italic=True)

# Taughannock Falls State Park
add_heading_styled('Taughannock Falls State Park', level=3)
sub_p = doc.add_paragraph()
add_run(sub_p, 'Three stories taller than Niagara, and you can walk right up to it', italic=True)
p = doc.add_paragraph()
add_run(p, 'Two miles north of ')
add_hyperlink(p, 'Trumansburg', 'https://www.google.com/maps/search/?api=1&query=Trumansburg+NY', MAPS_GREEN)
add_run(p, ' on Route 89. The 215-foot drop is the tallest single-drop waterfall east of the Rockies. Flat gorge trail to the base of the falls, under a mile. Swimming area and marina on the lake, summer concert series. The Black Diamond Trail runs from Taughannock south to the Cayuga Waterfront Trail in ')
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, ' \u2014 roughly 12 miles, flat, good for cycling. Go in April when the snowmelt is running and the sound reaches the trailhead.')

# Buttermilk Falls and Robert H. Treman
add_heading_styled('Buttermilk Falls and Robert H. Treman State Parks', level=3)
sub_p = doc.add_paragraph()
add_run(sub_p, 'The half-day gorge loop south of town', italic=True)
p = doc.add_paragraph()
add_run(p, 'Both within five miles south of downtown ')
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, '. ')
add_hyperlink(p, 'Buttermilk Falls', 'https://www.google.com/maps/search/?api=1&query=Buttermilk+Falls+State+Park+Ithaca+NY', MAPS_GREEN)
add_run(p, ' has a natural swimming pool at the base \u2014 cold, clear, the kind of swim people remember for years. ')
add_hyperlink(p, 'Robert H. Treman State Park', 'https://www.google.com/maps/search/?api=1&query=Robert+H+Treman+State+Park+Ithaca+NY', MAPS_GREEN)
add_run(p, '\'s Enfield Glen contains 12 waterfalls including the 115-foot Lucifer Falls, which is worth the hike by itself. Together they form the core of a half-day loop from ')
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, '.')

# Cascadilla Gorge
add_heading_styled('Cascadilla Gorge', level=3)
sub_p = doc.add_paragraph()
add_run(sub_p, 'For some Cornell students, this is just the commute', italic=True)
p = doc.add_paragraph()
add_run(p, 'Runs from downtown ')
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, ' up to the ')
add_hyperlink(p, 'Cornell', 'https://en.wikipedia.org/wiki/Cornell_University', WIKI_BLUE)
add_run(p, ' campus \u2014 400 feet of elevation through layered Devonian shale, six waterfalls, stone staircases built in the 1920s. The gorge trail is a geological classroom in motion: each layer of exposed rock tells a different chapter of the 380-million-year Devonian story.')

# Montezuma
add_heading_styled('Montezuma National Wildlife Refuge', level=3)
sub_p = doc.add_paragraph()
add_run(sub_p, 'A million ducks in a single autumn', italic=True)
p = doc.add_paragraph()
add_run(p, 'At the lake\'s north end. Over 300 bird species documented. The Cayuga people called the lake Tiohero \u2014 lake of flags or rushes \u2014 because of these northern marshes. Over a million ducks counted here in a single autumn. The auto tour loop is 4.5 miles through marsh and upland habitat. Fall and spring migration seasons are best; serious birders should come at dawn.')

add_image_placeholder('IMAGE: Montezuma National Wildlife Refuge \u2014 marsh and migratory birds, Cayuga Lake north end')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Montezuma National Wildlife Refuge', 'https://www.google.com/maps/search/?api=1&query=Montezuma+National+Wildlife+Refuge+NY', MAPS_GREEN)
add_run(cap, '. The Cayuga people called Cayuga Lake Tiohero \u2014 \'lake of flags or rushes\' \u2014 for these northern marshes.', italic=True)

# Named sites list
sites_p = doc.add_paragraph()
add_run(sites_p, 'Named sites: ', bold=True)
outdoor_sites = [
    ('Taughannock Falls State Park', 'https://www.google.com/maps/search/?api=1&query=Taughannock+Falls+State+Park+NY'),
    ('Buttermilk Falls State Park', 'https://www.google.com/maps/search/?api=1&query=Buttermilk+Falls+State+Park+Ithaca+NY'),
    ('Robert H. Treman State Park', 'https://www.google.com/maps/search/?api=1&query=Robert+H+Treman+State+Park+Ithaca+NY'),
    ('Cascadilla Gorge', 'https://www.google.com/maps/search/?api=1&query=Cascadilla+Gorge+Ithaca+NY'),
    ('Long Point State Park', 'https://www.google.com/maps/search/?api=1&query=Long+Point+State+Park+Aurora+NY'),
    ('Montezuma National Wildlife Refuge', 'https://www.google.com/maps/search/?api=1&query=Montezuma+National+Wildlife+Refuge+NY'),
]
for i, (name, url) in enumerate(outdoor_sites):
    add_hyperlink(sites_p, name, url, MAPS_GREEN)
    if i < len(outdoor_sites) - 1:
        add_run(sites_p, ' \u00b7 ')

# ============================================================
# WHERE TO STAY (preserved as-is, excluded from this pass)
# ============================================================
add_heading_styled('Where to stay', level=2)

stays = [
    ('Inns of Aurora', 'Aurora', 'Five historic properties in the National Historic District village, all within walking distance on Main Street. The Aurora Inn (1833, E.B. Morgan) is the flagship \u2014 ten rooms, shared balconies, the 1833 Kitchen & Bar downstairs. The EB Morgan House has working fireplaces and private lakefront access. Rowland House is decorated throughout in MacKenzie-Childs. The spa sits on the hill above the village. This is the full Cayuga Lake experience concentrated into one walkable village.', 'https://www.innsofaurora.com/', 'https://www.google.com/maps/search/?api=1&query=Inns+of+Aurora+NY'),
    ('Argos Inn', 'Ithaca', 'A LEED-certified 19th-century mansion a short walk from the Commons and Cascadilla Gorge. Boutique rooms, strong design character, the Argos Bar. The best independently owned option in Ithaca for people who want to be in the city, not outside it.', 'https://www.argosinn.com/', 'https://www.google.com/maps/search/?api=1&query=Argos+Inn+Ithaca+NY'),
    ('10 Fitch Boutique Inn', 'Auburn', 'Small boutique B&B in downtown Auburn, walking distance from Harriet Tubman NHP, Seward House, and the Equal Rights Heritage Center. The logical base for the heritage itinerary.', 'https://www.10fitch.com/', 'https://www.google.com/maps/search/?api=1&query=10+Fitch+Auburn+NY'),
    ('Springside Inn', 'Auburn', 'On the west shore of Owasco Lake south of Auburn. The innkeeper has led guests on walks through the woods behind the property to the former Underground Railroad path. The kind of detail that doesn\'t appear on the website but that guests who know to ask about remember.', 'https://www.springsideinn.com/', 'https://www.google.com/maps/search/?api=1&query=Springside+Inn+Auburn+NY'),
    ('The Evermore', 'King Ferry', 'The former MacKenzie-Childs estate, now a B&B. East-shore setting, distinctive design character, views of the lake. For travelers who want the lake experience in a quieter register than the Inns of Aurora.', 'https://TBD', 'https://www.google.com/maps/search/?api=1&query=The+Evermore+King+Ferry+NY'),
    ('Inn at Taughannock Falls', 'Trumansburg', 'Perched above the gorge at Taughannock Falls State Park. Dining with gorge views. The base for the west-shore outdoor day.', 'https://www.taughannock.com/', 'https://www.google.com/maps/search/?api=1&query=Inn+at+Taughannock+Falls+Trumansburg+NY'),
]

for name, loc, desc, official_url, maps_url in stays:
    p = doc.add_paragraph()
    add_hyperlink(p, name, official_url, OTHER_PURPLE)
    add_run(p, ' \u00b7 ')
    add_hyperlink(p, loc, maps_url, MAPS_GREEN)
    doc.add_paragraph(desc)

print("Part 5 done: outdoor, where to stay")

# ============================================================
# GETTING HERE AND GETTING AROUND
# ============================================================
add_heading_styled('Getting here and getting around', level=2)

p = doc.add_paragraph()
add_run(p, 'Cayuga Lake requires a car. ')
add_hyperlink(p, 'Ithaca Tompkins Regional Airport', 'https://www.google.com/maps/search/?api=1&query=Ithaca+Tompkins+Regional+Airport', MAPS_GREEN)
add_run(p, ' has limited routes. ')
add_hyperlink(p, 'Syracuse Hancock International Airport', 'https://www.google.com/maps/search/?api=1&query=Syracuse+Hancock+International+Airport', MAPS_GREEN)
add_run(p, ' is an hour north. Most visitors from New York City drive \u2014 four hours, Route 17 to Route 96.')

# Drive times table
add_heading_styled('Drive times on the lake', level=3)
drive_data = [
    ('Origin', 'Destination', 'Time'),
    ('Ithaca', 'Trumansburg', '20 minutes'),
    ('Ithaca', 'Aurora', '40 minutes'),
    ('Ithaca', 'Auburn', '35 minutes'),
    ('Auburn', 'Seneca Falls', '20 minutes'),
    ('Aurora', 'Auburn', '25 minutes'),
]
drive_table = doc.add_table(rows=len(drive_data), cols=3)
drive_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (origin, dest, time) in enumerate(drive_data):
    row = drive_table.rows[i]
    for j, val in enumerate([origin, dest, time]):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

# Shore roads
add_heading_styled('The two shore roads', level=3)
p = doc.add_paragraph()
add_run(p, 'Route 89 runs the west shore from ')
add_hyperlink(p, 'Ithaca', 'https://www.google.com/maps/search/?api=1&query=Ithaca+NY', MAPS_GREEN)
add_run(p, ' north through ')
add_hyperlink(p, 'Trumansburg', 'https://www.google.com/maps/search/?api=1&query=Trumansburg+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Interlaken', 'https://www.google.com/maps/search/?api=1&query=Interlaken+NY', MAPS_GREEN)
add_run(p, ', and ')
add_hyperlink(p, 'Romulus', 'https://www.google.com/maps/search/?api=1&query=Romulus+NY', MAPS_GREEN)
add_run(p, ' \u2014 primary wine trail loop, farm corridor, farm stands June through October. Route 90 runs the east shore \u2014 ')
add_hyperlink(p, 'Aurora', 'https://www.google.com/maps/search/?api=1&query=Aurora+NY+13026', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'King Ferry', 'https://www.google.com/maps/search/?api=1&query=King+Ferry+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Union Springs', 'https://www.google.com/maps/search/?api=1&query=Union+Springs+NY', MAPS_GREEN)
add_run(p, ' \u2014 quieter, more agricultural. The two roads connect at the lake\'s north end.')

# Seasonal notes
add_heading_styled('Seasonal notes', level=3)
p = doc.add_paragraph()
add_run(p, 'Wine trail tasting rooms: full hours May through November, reduced or closed December through April \u2014 call ahead. Farmers Market: Saturday April through December, Sunday May through November. Gorge trails best May through October; ')
add_hyperlink(p, 'Montezuma', 'https://www.google.com/maps/search/?api=1&query=Montezuma+National+Wildlife+Refuge+NY', MAPS_GREEN)
add_run(p, ' best in fall and spring migration.')

# ============================================================
# STAMP / BADGE CALLOUT (reformatted as distinct section)
# ============================================================
add_heading_styled('Stamp and badge callout', level=2)

# Wine trail stamps
wine_p = doc.add_paragraph()
add_run(wine_p, 'Wine trail stamps: ', bold=True)
add_run(wine_p, '21 producers tagged to the ')
add_hyperlink(wine_p, 'Cayuga Lake Wine Trail', 'https://upstate.tourismo.app/trails/cayuga-lake-wine-trail', UPSTATE_RED)
add_run(wine_p, ' badge. Count toward the Finger Lakes Wine master badge.')

# Heritage stamps
heritage_p = doc.add_paragraph()
add_run(heritage_p, 'Heritage stamps: ', bold=True)
heritage_sites_list = [
    ('Harriet Tubman NHP', 'https://www.nps.gov/hart/'),
    ('Seward House', 'https://www.sewardhouse.org/'),
    ('Fort Hill Cemetery', None),
    ('Willard Memorial Chapel', None),
    ("Women's Rights NHP", 'https://www.nps.gov/wori/'),
    ('Howland Stone Store', None),
    ('Montezuma Heritage Park', None),
]
for i, (name, url) in enumerate(heritage_sites_list):
    if url:
        add_hyperlink(heritage_p, name, url, OTHER_PURPLE)
    else:
        add_run(heritage_p, name)
    if i < len(heritage_sites_list) - 1:
        add_run(heritage_p, ', ')
add_run(heritage_p, ' \u2014 all ')
add_hyperlink(heritage_p, 'Finger Lakes Heritage Trail', 'https://upstate.tourismo.app/trails/finger-lakes-heritage-trail', UPSTATE_RED)
add_run(heritage_p, ' badge.')

# Nature stamps
nature_p = doc.add_paragraph()
add_run(nature_p, 'Nature stamps: ', bold=True)
nature_sites = [
    ('Taughannock Falls State Park', 'https://www.google.com/maps/search/?api=1&query=Taughannock+Falls+State+Park+NY'),
    ('Buttermilk Falls State Park', 'https://www.google.com/maps/search/?api=1&query=Buttermilk+Falls+State+Park+Ithaca+NY'),
    ('Robert H. Treman State Park', 'https://www.google.com/maps/search/?api=1&query=Robert+H+Treman+State+Park+Ithaca+NY'),
    ('Montezuma National Wildlife Refuge', 'https://www.google.com/maps/search/?api=1&query=Montezuma+National+Wildlife+Refuge+NY'),
]
for i, (name, url) in enumerate(nature_sites):
    add_hyperlink(nature_p, name, url, MAPS_GREEN)
    if i < len(nature_sites) - 1:
        add_run(nature_p, ', ')
add_run(nature_p, ' \u2014 Finger Lakes Outdoor badge.')

# Farm stamps
farm_p = doc.add_paragraph()
add_run(farm_p, 'Farm stamps: ', bold=True)
farm_sites = [
    ('Finger Lakes Cider House', 'https://www.fingerlakesciderhouse.com/'),
    ('Lively Run Dairy', 'https://www.livelyrun.com/'),
]
for i, (name, url) in enumerate(farm_sites):
    add_hyperlink(farm_p, name, url, OTHER_PURPLE)
    if i < len(farm_sites) - 1:
        add_run(farm_p, ', ')
add_run(farm_p, ', Route 89 farm stands \u2014 ')
add_hyperlink(farm_p, 'Finger Lakes Farm Trail', 'https://upstate.tourismo.app/trails/finger-lakes-farm-trail', UPSTATE_RED)
add_run(farm_p, ' badge.')

# ============================================================
# CTA BLOCK
# ============================================================
# Add a shaded CTA block with red border
cta_heading = doc.add_paragraph()
add_run(cta_heading, 'Explore more of the Finger Lakes.', bold=True, size=13)

# Apply shading to all CTA paragraphs
def apply_cta_shading(paragraph):
    """Apply light blue shading to paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{CTA_BG}"/>')
    pPr.append(shd)

apply_cta_shading(cta_heading)

cta_p1 = doc.add_paragraph()
add_run(cta_p1, 'Cayuga is the longest lake and the most historically layered. When you\'re ready to go deeper into the region, the ')
add_hyperlink(cta_p1, 'Seneca Lake chapter', 'https://upstate.tourismo.app/itineraries/seneca-lake', UPSTATE_RED)
add_run(cta_p1, ' picks up the wine story thirty-eight miles west, and the ')
add_hyperlink(cta_p1, 'Finger Lakes regional guide', 'https://upstatebound.com/guides/finger-lakes-region-04b4ec42-84de-4c38-aa0a-f689dc88d7a6', UPSTATE_RED)
add_run(cta_p1, ' has the full picture.')
apply_cta_shading(cta_p1)

cta_p2 = doc.add_paragraph()
add_hyperlink(cta_p2, 'Plan your visit on Upstate \u2192', 'https://upstate.tourismo.app/itineraries/cayuga-lake', UPSTATE_RED)
apply_cta_shading(cta_p2)

cta_p3 = doc.add_paragraph()
add_run(cta_p3, 'Related: ')
add_hyperlink(cta_p3, 'The Freedom Line Heritage Itinerary', 'https://upstate.tourismo.app/itineraries/the-freedom-line', UPSTATE_RED)
add_run(cta_p3, ' \u00b7 ')
add_hyperlink(cta_p3, 'The Cayuga Farm Loop', 'https://upstate.tourismo.app/itineraries/cayuga-farm-loop', UPSTATE_RED)
apply_cta_shading(cta_p3)

# ============================================================
# SAVE
# ============================================================
output_dir = '/mnt/user-data/outputs'
os.makedirs(output_dir, exist_ok=True)

output_path = os.path.join(output_dir, 'Cayuga_Lake_Chapter_v2.docx')
doc.save(output_path)
print(f"\nSaved to: {output_path}")

# Also save a copy to the repo
repo_path = '/home/user/fingerlakes/Cayuga_Lake_Chapter_v2.docx'
doc.save(repo_path)
print(f"Also saved to: {repo_path}")
print("\nDone! Chapter includes:")
print("  - TLDR ('The long and short of it') added before intro")
print("  - Intro restructured: 3 canonical paragraphs + Aurora connecting thread")
print("  - 'The reform corridor' renamed to 'The Heritage Corridor'")
print("  - 'Outside' renamed to 'Outdoor and nature'")
print("  - Color-coded hyperlinks throughout (Wiki blue, Maps green, Upstate red, Other purple)")
print("  - Stamp/badge callout reformatted as distinct section")
print("  - CTA block added at chapter end")
