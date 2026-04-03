#!/usr/bin/env python3
"""Generate the fixed Seneca Lake chapter as a properly formatted .docx"""

from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os

# --- Setup ---
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

for section in doc.sections:
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- Define Hyperlink character style ---
styles_element = doc.styles.element
hyperlink_style = parse_xml(
    f'<w:style {nsdecls("w")} w:type="character" w:styleId="Hyperlink">'
    f'  <w:name w:val="Hyperlink"/>'
    f'  <w:basedOn w:val="DefaultParagraphFont"/>'
    f'  <w:uiPriority w:val="99"/>'
    f'  <w:unhideWhenUsed/>'
    f'  <w:rPr>'
    f'    <w:color w:val="0563C1"/>'
    f'    <w:u w:val="single"/>'
    f'  </w:rPr>'
    f'</w:style>'
)
styles_element.append(hyperlink_style)

# --- Colors ---
WIKI_BLUE = '0563C1'
MAPS_GREEN = '1A7340'
UPSTATE_RED = 'C0392B'
OTHER_PURPLE = '6C3483'
CTA_BG = 'EAF4FB'

def add_hyperlink(paragraph, text, url, color_hex):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')
    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), color_hex)
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')
    rPr.append(sz)
    run_elem.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)
    return paragraph

def add_run(paragraph, text, bold=False, italic=False, size=None):
    run = paragraph.add_run(text)
    run.font.name = 'Arial'
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    return run

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
    return h

def add_image_placeholder(text):
    p = doc.add_paragraph()
    add_run(p, f'[ {text} ]', italic=True, size=10)
    return p

def apply_cta_shading(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{CTA_BG}"/>')
    pPr.append(shd)

# ============================================================
# CHAPTER HEADER
# ============================================================
header_p = doc.add_paragraph()
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(header_p, 'FINGER LAKES / SENECA LAKE', bold=True, size=10)

title = doc.add_heading('Seneca Lake', level=1)
for run in title.runs:
    run.font.name = 'Arial'

tagline = doc.add_paragraph()
add_run(tagline, 'The deepest lake entirely within New York State. The thermal engine that made the wine trail possible. Thirty-eight miles of shale slopes producing Riesling that earns the Mosel comparison it keeps being given.', italic=True)

meta = doc.add_paragraph()
add_run(meta, 'Chapter draft \u00b7 Schuyler, Seneca, Yates, and Ontario counties \u00b7 March 2026', size=10)

add_image_placeholder('IMAGE: Seneca Lake aerial, looking south toward Watkins Glen \u2014 fall preferred')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Seneca Lake', 'https://en.wikipedia.org/wiki/Seneca_Lake_(New_York)', WIKI_BLUE)
add_run(cap, ', looking south toward ', italic=True)
add_hyperlink(cap, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(cap, '. The deepest glacial lake entirely within New York State.', italic=True)

# ============================================================
# TLDR
# ============================================================
add_heading_styled('The long and short of it', level=3)
tldr = doc.add_paragraph()
add_run(tldr, 'Seneca is the deepest lake entirely within New York State, and its thermal mass is the reason the wine trail here has more producers than any other loop in the region. Run the west shore on day one and the east shore on day two: Route 414 is quieter, steeper, and produces wines that earn the Mosel comparison they\'re always being given. Get to ')
add_hyperlink(tldr, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(tldr, ' before nine in the morning or after four in the afternoon.')

# ============================================================
# INTRO
# ============================================================
add_heading_styled('The lake, and what it keeps producing', level=2)

p = doc.add_paragraph()
add_hyperlink(p, 'Seneca Lake', 'https://en.wikipedia.org/wiki/Seneca_Lake_(New_York)', WIKI_BLUE)
add_run(p, ' looks like the other lakes from the road, long, narrow, framed by vineyard slopes, but the numbers are different. At 618 feet deep, its bottom sits 173 feet below sea level, and it holds roughly half the water in all eleven ')
add_hyperlink(p, 'Finger Lakes', 'https://en.wikipedia.org/wiki/Finger_Lakes', WIKI_BLUE)
add_run(p, ' combined. It freezes completely about once a century. The last time was 1912, when people skated the full 35 miles from ')
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, ' to ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, '. That thermal mass is the whole story. The lake holds summer heat deep into autumn, keeps the surrounding slopes from freezing when the plateau above is already locked in, and creates a microclimate that makes viticulture possible at this latitude.')

p2 = doc.add_paragraph()
add_run(p2, 'The ')
add_hyperlink(p2, 'Seneca Nation', 'https://en.wikipedia.org/wiki/Seneca_Nation_of_Indians', WIKI_BLUE)
add_run(p2, ' farmed these shores and built their capital, Kanadeseaga, near what is now ')
add_hyperlink(p2, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p2, ', before Washington\'s troops burned their villages in 1779. Catherine\'s Town, near present-day ')
add_hyperlink(p2, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p2, ', was among those destroyed. The early nursery industry that made Geneva\'s founding wealth began with the Kanadesaga fruit trees that survived Washington\'s campaign. ')
add_hyperlink(p2, 'Hobart and William Smith Colleges', 'https://www.google.com/maps/search/?api=1&query=Hobart+and+William+Smith+Colleges+Geneva+NY', MAPS_GREEN)
add_run(p2, ' trace their origins to Geneva Academy, established 1797 at the lake\'s north end, on land that was ')
add_hyperlink(p2, 'Seneca Nation', 'https://en.wikipedia.org/wiki/Seneca_Nation_of_Indians', WIKI_BLUE)
add_run(p2, ' territory. The institution\'s specific claim: ')
add_hyperlink(p2, 'Elizabeth Blackwell', 'https://en.wikipedia.org/wiki/Elizabeth_Blackwell', WIKI_BLUE)
add_run(p2, ' applied to and was rejected by 29 medical schools before being admitted to Geneva Medical College in 1847. She graduated two years later at the head of her class, the first woman to receive a medical degree in the United States. This happened six miles west of ')
add_hyperlink(p2, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(p2, ', in the same decade as the ')
add_hyperlink(p2, "women's rights convention", 'https://en.wikipedia.org/wiki/Women%27s_Rights_National_Historical_Park', WIKI_BLUE)
add_run(p2, '.')

p3 = doc.add_paragraph()
add_run(p3, 'By 1900 there were more than 20,000 acres of vines across the Finger Lakes, with Seneca at the center. Prohibition dismantled most of it. Scientists at Cornell\'s Agricultural Experiment Station in ')
add_hyperlink(p3, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p3, ' spent subsequent decades convinced that European wine grapes could not survive these winters. In 1951, a Ukrainian viticulturalist named ')
add_hyperlink(p3, 'Konstantin Frank', 'https://en.wikipedia.org/wiki/Konstantin_Frank', WIKI_BLUE)
add_run(p3, ' looked at the cold Seneca slopes and recognized the conditions from home. He was right. ')
add_hyperlink(p3, 'Hermann Wiemer', 'https://en.wikipedia.org/wiki/Hermann_J._Wiemer_Vineyard', WIKI_BLUE)
add_run(p3, ', who came from Germany\'s Mosel Valley, saw the east-shore slopes in the 1970s and recognized the geology. He stayed for the rest of his working life. The Rieslings from those slopes have been earning the comparison ever since.')

p4 = doc.add_paragraph()
add_hyperlink(p4, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p4, ' anchors the north with a real downtown, ')
add_hyperlink(p4, 'Hobart and William Smith Colleges', 'https://www.google.com/maps/search/?api=1&query=Hobart+and+William+Smith+Colleges+Geneva+NY', MAPS_GREEN)
add_run(p4, ' on the lakefront, and the Cornell Experiment Station where the wine region\'s modern history was made. ')
add_hyperlink(p4, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p4, ' anchors the south with the best gorge trail in the Finger Lakes and a village that has been drawing visitors since the 1840s. Between them: 33 wineries, two shore roads, and the steepest east-facing shale slopes in the region.')

print("Part 1 done: header through intro")

# ============================================================
# AT A GLANCE
# ============================================================
add_heading_styled('At a glance', level=2)
glance_data = [
    ('Length', '38 miles'),
    ('Depth', '618 feet'),
    ('Counties', 'Schuyler, Seneca, Yates, Ontario'),
    ('Wine trail', '33 producers'),
    ('From NYC', '4 hours'),
    ('From Buffalo', '2 hours'),
    ('From Syracuse', '1 hour'),
    ('Anchor cities', 'Geneva \u00b7 Watkins Glen'),
]
table = doc.add_table(rows=len(glance_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (label, value) in enumerate(glance_data):
    row = table.rows[i]
    for j, val in enumerate([label, value]):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        if j == 0:
            run.bold = True

# ============================================================
# THE TOWNS
# ============================================================
add_heading_styled('The towns', level=2)

p = doc.add_paragraph()
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, ' anchors the north with colleges and a downtown that has functioned as a downtown since before the Erie Canal was built through it in 1825. ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, ' anchors the south with a gorge and a racing history it carries without strain. The east-shore towns, ')
add_hyperlink(p, 'Burdett', 'https://www.google.com/maps/search/?api=1&query=Burdett+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Hector', 'https://www.google.com/maps/search/?api=1&query=Hector+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Lodi', 'https://www.google.com/maps/search/?api=1&query=Lodi+NY', MAPS_GREEN)
add_run(p, ', are wine country in its working form. ')
add_hyperlink(p, 'Dundee', 'https://www.google.com/maps/search/?api=1&query=Dundee+NY', MAPS_GREEN)
add_run(p, ' sits mid-lake on the west shore with a benchmark winery that belongs on any serious list.')

add_image_placeholder('IMAGE: Geneva waterfront and Hobart and William Smith Colleges, Seneca Lake visible, summer or fall')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(cap, ', north end of Seneca Lake. ', italic=True)
add_hyperlink(cap, 'Hobart and William Smith Colleges', 'https://www.google.com/maps/search/?api=1&query=Hobart+and+William+Smith+Colleges+Geneva+NY', MAPS_GREEN)
add_run(cap, ' on the lakefront bluff.', italic=True)

# --- Geneva ---
add_heading_styled('Geneva', level=3)
meta = doc.add_paragraph()
add_run(meta, 'ONTARIO COUNTY \u00b7 NORTH END \u00b7 POP. ~12,500', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR THE NORTH-END BASE AND A PROPER DOWNTOWN', bold=True, size=10)

p = doc.add_paragraph()
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, ' has been at the lake\'s north end long enough to accumulate a biography most visitors miss. ')
add_hyperlink(p, 'Hobart and William Smith Colleges', 'https://www.google.com/maps/search/?api=1&query=Hobart+and+William+Smith+Colleges+Geneva+NY', MAPS_GREEN)
add_run(p, ' trace their origins to Geneva Academy, established 1797. The institution\'s specific claim: ')
add_hyperlink(p, 'Elizabeth Blackwell', 'https://en.wikipedia.org/wiki/Elizabeth_Blackwell', WIKI_BLUE)
add_run(p, ' applied to and was rejected by 29 medical schools before being admitted to Geneva Medical College in 1847. She graduated at the head of her class, the first woman to receive a medical degree in the United States. The campus sits on the lakefront bluff. The lake view from it is the best free view at the north end.')

p = doc.add_paragraph()
add_run(p, 'The Erie Canal bypassed ')
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, ', as did the railroads later, which is why it never became a major center of commerce, and why it has a downtown that still functions as one. The ')
add_hyperlink(p, 'Smith Opera House', 'https://www.google.com/maps/search/?api=1&query=Smith+Opera+House+Geneva+NY', MAPS_GREEN)
add_run(p, ' on Seneca Street has been presenting performances since 1894. Cornell\'s New York State Agricultural Experiment Station is here, the same institution where scientists spent decades insisting European wine grapes could not survive the Finger Lakes winters, and where ')
add_hyperlink(p, 'Konstantin Frank', 'https://en.wikipedia.org/wiki/Konstantin_Frank', WIKI_BLUE)
add_run(p, ' eventually proved them wrong. ')
add_hyperlink(p, 'Geneva on the Lake', 'https://www.google.com/maps/search/?api=1&query=Geneva+on+the+Lake+NY', MAPS_GREEN)
add_run(p, ', a restored 1911 lakefront villa, is the most distinctive lodging on Seneca.')

# --- Watkins Glen ---
add_image_placeholder('IMAGE: Watkins Glen State Park gorge trail, 19 waterfalls, stone walls, layered shale, summer preferred')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Watkins Glen State Park', 'https://en.wikipedia.org/wiki/Watkins_Glen_State_Park', WIKI_BLUE)
add_run(cap, '. The gorge trail winds 1.5 miles through 19 waterfalls past 200-foot cliffs.', italic=True)

add_heading_styled('Watkins Glen', level=3)
meta = doc.add_paragraph()
add_run(meta, 'SCHUYLER COUNTY \u00b7 SOUTH END \u00b7 POP. ~1,800', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR THE GORGE AND THE SOUTH-END BASE', bold=True, size=10)

p = doc.add_paragraph()
add_hyperlink(p, 'Watkins Glen State Park', 'https://en.wikipedia.org/wiki/Watkins_Glen_State_Park', WIKI_BLUE)
add_run(p, ' generates 19 waterfalls within two miles as Glen Creek descends 400 feet through 200-foot shale cliffs. The gorge trail winds through tunnels, over and under the falls, past water-carved shale walls. Get there before 9am or after 4pm in summer. The gorge fills up.')

p = doc.add_paragraph()
add_run(p, 'The village hosted its first road race in 1948 on a course that wrapped around the park and crossed Glen Creek. ')
add_hyperlink(p, 'Watkins Glen International', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+International+NY', MAPS_GREEN)
add_run(p, ' hosted Formula 1\'s US Grand Prix from 1961 to 1980 and still draws large racing crowds for NASCAR and vintage events. Then it goes back to being a small lakeside village with a walkable main street, good restaurants, and the gorge entrance two blocks from the waterfront. The ')
add_hyperlink(p, 'Harbor Hotel', 'https://www.google.com/maps/search/?api=1&query=Harbor+Hotel+Watkins+Glen+NY', MAPS_GREEN)
add_run(p, ' sits at the lake\'s edge. The wine trail runs north in both directions from here.')

# --- Dundee ---
add_heading_styled('Dundee', level=3)
meta = doc.add_paragraph()
add_run(meta, 'YATES COUNTY \u00b7 WEST SHORE, MID-LAKE \u00b7 POP. ~1,700', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR WEST-SHORE WINE WITH FEWER CROWDS', bold=True, size=10)

p = doc.add_paragraph()
add_hyperlink(p, 'Dundee', 'https://www.google.com/maps/search/?api=1&query=Dundee+NY', MAPS_GREEN)
add_run(p, ' sits on the west shore roughly mid-lake, quieter than the ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, ' corridor below and less trafficked than the ')
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, ' approaches above. ')
add_hyperlink(p, 'Hermann J. Wiemer Vineyard', 'https://en.wikipedia.org/wiki/Hermann_J._Wiemer_Vineyard', WIKI_BLUE)
add_run(p, ' is here, one of the benchmark producers on the entire trail, recognized by Wine & Spirits Magazine as one of the world\'s Top 100 Wineries. ')
add_hyperlink(p, 'Glenora Wine Cellars', 'https://www.google.com/maps/search/?api=1&query=Glenora+Wine+Cellars+Dundee+NY', MAPS_GREEN)
add_run(p, ', the first winery established on Seneca Lake, is a mile south of the village with an inn and restaurant overlooking the water. ')
add_hyperlink(p, 'Dundee', 'https://www.google.com/maps/search/?api=1&query=Dundee+NY', MAPS_GREEN)
add_run(p, ' itself is a small agricultural town without wine-trail branding. That is the reason to stop.')

# --- Burdett ---
add_heading_styled('Burdett', level=3)
meta = doc.add_paragraph()
add_run(meta, 'SCHUYLER COUNTY \u00b7 EAST SHORE, SOUTH SECTOR \u00b7 POP. ~350', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR THE EAST-SHORE ENTRY POINT', bold=True, size=10)

p = doc.add_paragraph()
add_hyperlink(p, 'Burdett', 'https://www.google.com/maps/search/?api=1&query=Burdett+NY', MAPS_GREEN)
add_run(p, ' is perched on the east-shore bluff where the vineyards begin in earnest. Hector Falls drops into view from the road, a pull-over waterfall visible from Route 414, with the lower cascades visible from the lake. ')
add_hyperlink(p, 'Atwater Vineyards', 'https://www.google.com/maps/search/?api=1&query=Atwater+Vineyards+Burdett+NY', MAPS_GREEN)
add_run(p, ' is here, with lake views and a reliable Riesling program. ')
add_hyperlink(p, 'Two Goats Brewing', 'https://www.google.com/maps/search/?api=1&query=Two+Goats+Brewing+Burdett+NY', MAPS_GREEN)
add_run(p, ' and ')
add_hyperlink(p, 'Finger Lakes Distilling', 'https://www.google.com/maps/search/?api=1&query=Finger+Lakes+Distilling+Burdett+NY', MAPS_GREEN)
add_run(p, ' are both in ')
add_hyperlink(p, 'Burdett', 'https://www.google.com/maps/search/?api=1&query=Burdett+NY', MAPS_GREEN)
add_run(p, ', making it the most complete stop for non-wine drinkers on the east shore. ')
add_hyperlink(p, 'Hillick & Hobbs Estate', 'https://www.google.com/maps/search/?api=1&query=Hillick+and+Hobbs+Estate+Burdett+NY', MAPS_GREEN)
add_run(p, ' produces some of the most precise small-production Riesling on this side of the lake.')

# --- Hector ---
add_heading_styled('Hector', level=3)
meta = doc.add_paragraph()
add_run(meta, 'SCHUYLER COUNTY \u00b7 EAST SHORE, MID-LAKE \u00b7 POP. ~4,500 TOWNSHIP', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR MID-EAST-SHORE STOPS AND THE NATIONAL FOREST', bold=True, size=10)

p = doc.add_paragraph()
add_hyperlink(p, 'Hector', 'https://www.google.com/maps/search/?api=1&query=Hector+NY', MAPS_GREEN)
add_run(p, ' is a township rather than a village, a wide stretch of the east shore between ')
add_hyperlink(p, 'Burdett', 'https://www.google.com/maps/search/?api=1&query=Burdett+NY', MAPS_GREEN)
add_run(p, ' to the south and ')
add_hyperlink(p, 'Lodi', 'https://www.google.com/maps/search/?api=1&query=Lodi+NY', MAPS_GREEN)
add_run(p, ' to the north, with vineyards visible on nearly every hillside. ')
add_hyperlink(p, 'Red Newt Cellars', 'https://www.google.com/maps/search/?api=1&query=Red+Newt+Cellars+Hector+NY', MAPS_GREEN)
add_run(p, ' is here, with a bistro that makes it a legitimate lunch stop as well as a tasting room. Hazlitt 1852 Vineyards has been farming this land since its name implies. The ')
add_hyperlink(p, 'Finger Lakes National Forest', 'https://www.google.com/maps/search/?api=1&query=Finger+Lakes+National+Forest+Hector+NY', MAPS_GREEN)
add_run(p, ' headquarters is on Route 414 in ')
add_hyperlink(p, 'Hector', 'https://www.google.com/maps/search/?api=1&query=Hector+NY', MAPS_GREEN)
add_run(p, ', the entry point for 16,000 acres of trails between Seneca and Cayuga Lakes.')

add_image_placeholder('IMAGE: Finger Lakes National Forest, pastureland and vineyards on the ridge between Seneca and Cayuga, summer or fall')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Finger Lakes National Forest', 'https://www.google.com/maps/search/?api=1&query=Finger+Lakes+National+Forest+Hector+NY', MAPS_GREEN)
add_run(cap, ', Hector. New York\'s only national forest. The pastures are still leased to working farms.', italic=True)

# --- Lodi ---
add_heading_styled('Lodi', level=3)
meta = doc.add_paragraph()
add_run(meta, 'SENECA COUNTY \u00b7 EAST SHORE, UPPER MID-LAKE \u00b7 POP. ~600', bold=True, size=10)
best = doc.add_paragraph()
add_run(best, 'BEST FOR THE EAST-SHORE ANCHOR STOP', bold=True, size=10)

p = doc.add_paragraph()
add_hyperlink(p, 'Lodi', 'https://www.google.com/maps/search/?api=1&query=Lodi+NY', MAPS_GREEN)
add_run(p, ' is where the east shore concentrates. ')
add_hyperlink(p, 'Wagner Vineyards', 'https://www.google.com/maps/search/?api=1&query=Wagner+Vineyards+Lodi+NY', MAPS_GREEN)
add_run(p, ' is the most complete operation on Seneca, estate-grown winery, brewery, and restaurant under one roof, with panoramic lake views. ')
add_hyperlink(p, 'Boundary Breaks', 'https://www.google.com/maps/search/?api=1&query=Boundary+Breaks+Vineyard+Lodi+NY', MAPS_GREEN)
add_run(p, ' is a single-vineyard Riesling program. ')
add_hyperlink(p, 'Lamoreaux Landing', 'https://www.google.com/maps/search/?api=1&query=Lamoreaux+Landing+Wine+Cellars+Lodi+NY', MAPS_GREEN)
add_run(p, ' has dramatic lakefront architecture and one of the more serious east-shore programs. ')
add_hyperlink(p, 'Silver Thread Vineyard', 'https://www.google.com/maps/search/?api=1&query=Silver+Thread+Vineyard+Lodi+NY', MAPS_GREEN)
add_run(p, ' is a small organic producer with a reputation that outpaces its size. If you are only making one east-shore stop, ')
add_hyperlink(p, 'Lodi', 'https://www.google.com/maps/search/?api=1&query=Lodi+NY', MAPS_GREEN)
add_run(p, ' is where to make it.')

print("Part 2 done: at-a-glance and towns")

# ============================================================
# THE HERITAGE CORRIDOR
# ============================================================
add_heading_styled('The Heritage Corridor', level=2)

add_image_placeholder('IMAGE: Elizabeth Blackwell portrait or Geneva Medical College historical photo, public domain')
cap = doc.add_paragraph()
add_hyperlink(cap, 'Elizabeth Blackwell', 'https://en.wikipedia.org/wiki/Elizabeth_Blackwell', WIKI_BLUE)
add_run(cap, ', first woman to receive a medical degree in the United States, graduated from Geneva Medical College in 1849.', italic=True)

p = doc.add_paragraph()
add_run(p, 'Seneca County carries more of the reform era than most visitors realize. ')
add_hyperlink(p, 'Waterloo', 'https://www.google.com/maps/search/?api=1&query=Waterloo+NY', MAPS_GREEN)
add_run(p, ', five miles northwest of ')
add_hyperlink(p, 'Seneca Falls', 'https://www.google.com/maps/search/?api=1&query=Seneca+Falls+NY', MAPS_GREEN)
add_run(p, ', holds two pieces of American history on the same main street: the ')
add_hyperlink(p, "M'Clintock House", 'https://www.google.com/maps/search/?api=1&query=MClintock+House+Waterloo+NY', MAPS_GREEN)
add_run(p, ', where the Declaration of Sentiments was drafted before the 1848 convention, and the birthplace of Memorial Day. On May 5, 1866, pharmacist Henry Welles and General John Murray organized the first community-wide observance of Memorial Day in ')
add_hyperlink(p, 'Waterloo', 'https://www.google.com/maps/search/?api=1&query=Waterloo+NY', MAPS_GREEN)
add_run(p, '. In 1966, President Johnson signed a proclamation formally recognizing it.')

p = doc.add_paragraph()
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, '\'s claim is ')
add_hyperlink(p, 'Elizabeth Blackwell', 'https://en.wikipedia.org/wiki/Elizabeth_Blackwell', WIKI_BLUE)
add_run(p, '. The reform movements of the 19th century were not confined to the Auburn\u2013Seneca Falls corridor. They ran the length of the lake.')

# Heritage sites list
add_heading_styled('Heritage sites, full list', level=3)
badge = doc.add_paragraph()
add_run(badge, 'Badge track: All sites below are tagged to the ')
add_hyperlink(badge, 'Finger Lakes Heritage Trail', 'https://upstate.tourismo.app/trails/finger-lakes-heritage-trail', UPSTATE_RED)
add_run(badge, ' passport badge.')

sites = [
    ("Women's Rights National Historical Park", 'Seneca Falls', 'Where 300 people rewrote the rules in two days', 'The Wesleyan Chapel where the 1848 convention met, the Stanton House, and the M\'Clintock House in Waterloo where the Declaration of Sentiments was drafted. Start at the NPS visitor center on Fall Street. The chapel interior has been partially reconstructed; the original walls are visible behind glass. Free admission.', 'https://www.nps.gov/wori/', 'https://www.google.com/maps/search/?api=1&query=Womens+Rights+National+Historical+Park+Seneca+Falls+NY'),
    ("M'Clintock House", 'Waterloo NY', 'The kitchen table where the Declaration was written', 'Where Elizabeth Cady Stanton and her collaborators drafted the Declaration of Sentiments in 1848, days before the convention. Part of Women\'s Rights NHP. The house is restored to its 1848 appearance. Open seasonally; check NPS schedule.', None, 'https://www.google.com/maps/search/?api=1&query=MClintock+House+Waterloo+NY'),
    ('National Memorial Day Museum', 'Waterloo NY', 'The first observance, before it had a name', '1866. Pharmacist Henry Welles and General John Murray organized the first community-wide Memorial Day observance in Waterloo. Congressional recognition came a century later in 1966. The museum documents the origins and the town\'s claim. Same main street as the M\'Clintock House.', None, 'https://www.google.com/maps/search/?api=1&query=National+Memorial+Day+Museum+Waterloo+NY'),
    ('New York State Agricultural Experiment Station', 'Geneva NY', 'Where the experts were wrong and Konstantin Frank was right', 'Cornell\'s research station where scientists spent decades insisting European wine grapes could not survive Finger Lakes winters. In 1951, Konstantin Frank arrived and proved them wrong. The station\'s work on cold-climate viticulture continues. Grounds open to visitors; the research plots are visible from the road.', None, 'https://www.google.com/maps/search/?api=1&query=NY+Agricultural+Experiment+Station+Geneva+NY'),
    ('Hobart and William Smith Colleges', 'Geneva NY', 'Known for Elizabeth Blackwell and the best free lake view at the north end', 'The institution where Elizabeth Blackwell became the first woman to receive a medical degree in the United States in 1849, after being rejected by 29 other schools. The lakefront campus sits on a bluff above Seneca Lake. The view from the campus is the best free view at the north end. Open to visitors.', None, 'https://www.google.com/maps/search/?api=1&query=Hobart+and+William+Smith+Colleges+Geneva+NY'),
]
for name, loc, subtitle, desc, official_url, maps_url in sites:
    p = doc.add_paragraph()
    add_run(p, '\u2014 ')
    if official_url:
        add_hyperlink(p, name, official_url, OTHER_PURPLE)
    else:
        add_run(p, name, bold=True)
    add_run(p, ' ')
    add_hyperlink(p, loc, maps_url, MAPS_GREEN)
    sub_p = doc.add_paragraph()
    add_run(sub_p, subtitle, italic=True)
    doc.add_paragraph(desc)

print("Part 3 done: heritage corridor")

# ============================================================
# THE SENECA LAKE WINE TRAIL
# ============================================================
add_heading_styled('The Seneca Lake Wine Trail', level=2)

add_image_placeholder('IMAGE: Hermann J. Wiemer Vineyard, east shore slopes, Seneca Lake visible below, fall preferred')
cap = doc.add_paragraph()
add_run(cap, 'East shore, Seneca Lake. The shale slopes facing west produce the most precise Rieslings on the trail.', italic=True)

p = doc.add_paragraph()
add_run(p, 'The ')
add_hyperlink(p, 'Seneca Lake Wine Trail', 'https://upstate.tourismo.app/trails/seneca-lake-wine-trail', UPSTATE_RED)
add_run(p, ' has 33 producers. What distinguishes Seneca: more thermal mass, lower elevation on the lake shores, steeper east-side slopes, and a slightly richer, rounder expression in the Riesling. The east shore, Route 414 produces wines that earn serious Mosel comparisons because the conditions are genuinely similar. Shale slopes, cold lake moderating temperatures, long hang time in autumn. The east shore\'s so-called banana belt, the southeast-facing slopes where afternoon sun lingers into the evening, produces the most concentrated fruit on the trail. ')
add_hyperlink(p, 'Hermann Wiemer', 'https://en.wikipedia.org/wiki/Hermann_J._Wiemer_Vineyard', WIKI_BLUE)
add_run(p, ' identified this in the 1970s. The wines from those slopes have been making the argument ever since.')

# How to run the trail
add_heading_styled('How to run the trail', level=3)
p = doc.add_paragraph()
add_run(p, 'West shore: Route 14 north from ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, ' through ')
add_hyperlink(p, 'Rock Stream', 'https://www.google.com/maps/search/?api=1&query=Rock+Stream+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Dundee', 'https://www.google.com/maps/search/?api=1&query=Dundee+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Himrod', 'https://www.google.com/maps/search/?api=1&query=Himrod+NY', MAPS_GREEN)
add_run(p, ', to ')
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, '. More producers, more established names, and easier logistics. East shore: Route 414 north from ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, ' through ')
add_hyperlink(p, 'Burdett', 'https://www.google.com/maps/search/?api=1&query=Burdett+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Hector', 'https://www.google.com/maps/search/?api=1&query=Hector+NY', MAPS_GREEN)
add_run(p, ', and ')
add_hyperlink(p, 'Lodi', 'https://www.google.com/maps/search/?api=1&query=Lodi+NY', MAPS_GREEN)
add_run(p, '. Quieter, steeper, the better second day when you want depth over breadth. The two roads connect at ')
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, ' on the north end and ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, ' on the south, so a two-day loop covering both shores is straightforward: one road each day, no backtracking. Tasting rooms run full hours May through November; call ahead in shoulder season.')

# Producers worth naming
add_heading_styled('Producers worth naming', level=3)
producers = [
    ('Hermann J. Wiemer Vineyard', 'Dundee NY \u00b7 west shore', 'Where the Mosel meets Seneca', 'The benchmark for what Seneca can do at its most precise. World\'s Top 100 Wineries. Biodynamic certified. Hermann Wiemer came from Germany\'s Mosel Valley, saw the west-shore shale slopes in the 1970s, and stayed for the rest of his working life.', 'https://www.google.com/maps/search/?api=1&query=Hermann+J+Wiemer+Vineyard+Dundee+NY'),
    ('Wagner Vineyards', 'Lodi NY \u00b7 east shore', 'Wine, beer, and lunch on the same hilltop', 'The most complete stop on the east shore. Estate-grown winery, brewery, and restaurant under one roof, with panoramic lake views. The Ginny Lee Cafe serves lunch on a terrace above the water. Open daily year-round.', 'https://www.google.com/maps/search/?api=1&query=Wagner+Vineyards+Lodi+NY'),
    ('Boundary Breaks Vineyard', 'Lodi NY \u00b7 east shore', 'One vineyard, one variety, no buses', 'Single-vineyard Riesling program. Wine Enthusiast named the Dry Riesling to its Top 100 Wines in the World twice. No groups larger than six. The argument for why site designations matter in this AVA.', 'https://www.google.com/maps/search/?api=1&query=Boundary+Breaks+Vineyard+Lodi+NY'),
    ('Lamoreaux Landing Wine Cellars', 'Lodi NY \u00b7 east shore', 'Three generations and the best architecture on the trail', 'Dramatic Greek Revival tasting room with an unobstructed lake view. The Lamoreaux family has been farming grapes at this Lodi hillside for three generations. One of the more serious east-shore Riesling programs.', 'https://www.google.com/maps/search/?api=1&query=Lamoreaux+Landing+Wine+Cellars+Lodi+NY'),
    ('Glenora Wine Cellars', 'Dundee NY \u00b7 west shore', 'First on the lake, still the mid-shore overnight', 'The first winery on Seneca Lake, established 1977. The 30-room inn sits above the vineyard. Veraisons Restaurant changes its menu each season. The logical west-shore overnight.', 'https://www.google.com/maps/search/?api=1&query=Glenora+Wine+Cellars+Dundee+NY'),
    ('Red Newt Cellars', 'Hector NY \u00b7 east shore', 'The reason to take Route 414', 'Winery and farm-to-table bistro. The best lunch stop on the east shore. Thursday through Saturday for dinner; Sunday for the Jazz Brunch. Scott Signori opened it in 1999; his team has kept the same specificity since his passing in 2021.', 'https://www.google.com/maps/search/?api=1&query=Red+Newt+Cellars+Hector+NY'),
    ('Atwater Vineyards', 'Burdett NY \u00b7 east shore', 'The first stop heading north on Route 414', 'Consistent quality, lake views from the terrace, and enough range to orient first-time visitors. Close enough to Hillick & Hobbs and Two Goats that the first two hours on Route 414 sort themselves out.', 'https://www.google.com/maps/search/?api=1&query=Atwater+Vineyards+Burdett+NY'),
    ('Hillick & Hobbs Estate', 'Burdett NY \u00b7 east shore', 'Small production, longer conversations', 'Single east-shore vineyard, limited production, and a tasting experience that doesn\'t feel like a line. The estate farms the shale slopes that give east-shore Seneca its claim to Mosel comparison.', 'https://www.google.com/maps/search/?api=1&query=Hillick+and+Hobbs+Estate+Burdett+NY'),
]
for name, loc, subtitle, desc, maps_url in producers:
    p = doc.add_paragraph()
    add_run(p, '\u2014 ')
    add_hyperlink(p, name, maps_url, MAPS_GREEN)
    add_run(p, ' ' + loc)
    sub_p = doc.add_paragraph()
    add_run(sub_p, subtitle, italic=True)
    doc.add_paragraph(desc)

# Tier 1 producers
add_heading_styled('Tier 1', level=3)

# Hermann J. Wiemer
t1 = doc.add_paragraph()
add_run(t1, 'Hermann J. Wiemer Vineyard', bold=True)
sub = doc.add_paragraph()
add_run(sub, 'Where the Mosel meets Seneca', italic=True)
p = doc.add_paragraph()
add_hyperlink(p, 'Hermann Wiemer', 'https://en.wikipedia.org/wiki/Hermann_J._Wiemer_Vineyard', WIKI_BLUE)
add_run(p, ' grew up in the Mosel Valley, where his father grafted vines at the Geisenheim Institute and his family had been making wine for three centuries. When he arrived at Seneca Lake in 1968, he looked at the west-shore shale slopes outside ')
add_hyperlink(p, 'Dundee', 'https://www.google.com/maps/search/?api=1&query=Dundee+NY', MAPS_GREEN)
add_run(p, ' and recognized the geology from home. He planted vinifera on an abandoned soybean farm in 1973 against the advice of every expert in the region. By 1988 his Riesling was the first New York wine served in first class on an international flight. Wine & Spirits Magazine has named the estate one of the world\'s Top 100 Wineries multiple years running. Fred Merwarth, who apprenticed under Wiemer from 2001, took over in 2007 alongside Oskar Bynke and has since converted the 80-acre estate to biodynamic certification. The wines are dry and mineral-driven, made with indigenous yeasts from hand-sorted fruit. Open daily on Route 14 in ')
add_hyperlink(p, 'Dundee', 'https://www.google.com/maps/search/?api=1&query=Dundee+NY', MAPS_GREEN)
add_run(p, '.')

# Wagner Vineyards
t1 = doc.add_paragraph()
add_run(t1, 'Wagner Vineyards', bold=True)
sub = doc.add_paragraph()
add_run(sub, 'Wine, beer, and lunch on the same hilltop', italic=True)
p = doc.add_paragraph()
add_hyperlink(p, 'Wagner Vineyards', 'https://www.google.com/maps/search/?api=1&query=Wagner+Vineyards+Lodi+NY', MAPS_GREEN)
add_run(p, ' has been farming the east shore outside ')
add_hyperlink(p, 'Lodi', 'https://www.google.com/maps/search/?api=1&query=Lodi+NY', MAPS_GREEN)
add_run(p, ' long enough that the hamlet and the winery are effectively inseparable. The estate grows its own grapes, crushes on site, and runs Wagner Valley Brewing Company from the same hilltop, making it the most complete single stop on the ')
add_hyperlink(p, 'Seneca Lake Wine Trail', 'https://upstate.tourismo.app/trails/seneca-lake-wine-trail', UPSTATE_RED)
add_run(p, '. The tasting room and brewery share a building with the Ginny Lee Cafe, which serves lunch with vineyard and lake views from a terrace above the water. The portfolio spans Riesling, Chardonnay, Cabernet Franc, and red blends alongside the rotating brewery lineup. Open daily year-round. Watch for the signs on Route 414.')

# Red Newt Cellars
t1 = doc.add_paragraph()
add_run(t1, 'Red Newt Cellars', bold=True)
sub = doc.add_paragraph()
add_run(sub, 'The reason to take Route 414', italic=True)
p = doc.add_paragraph()
add_run(p, 'Scott Signori opened ')
add_hyperlink(p, 'Red Newt', 'https://www.google.com/maps/search/?api=1&query=Red+Newt+Cellars+Hector+NY', MAPS_GREEN)
add_run(p, ' in 1999 in a roadside building in ')
add_hyperlink(p, 'Hector', 'https://www.google.com/maps/search/?api=1&query=Hector+NY', MAPS_GREEN)
add_run(p, ' that had been the Bond Fruit and Dairy Stand for generations before him. The bistro runs alongside the tasting room, sources from local farms, and changes the menu with what is available. One of the few stops on the east shore worth a full afternoon rather than a quick pour and move on. The estate Riesling comes from hillside vineyards above the lake. Thursday through Saturday for dinner; Sunday for the Jazz Brunch. Scott passed away in 2021. The Signori family and his team have kept the same specificity he brought to it from the start.')

# Glenora Wine Cellars
t1 = doc.add_paragraph()
add_run(t1, 'Glenora Wine Cellars', bold=True)
sub = doc.add_paragraph()
add_run(sub, 'First on the lake, still the mid-shore overnight', italic=True)
p = doc.add_paragraph()
add_hyperlink(p, 'Glenora', 'https://www.google.com/maps/search/?api=1&query=Glenora+Wine+Cellars+Dundee+NY', MAPS_GREEN)
add_run(p, ' opened in 1977, the first winery established on Seneca Lake, on the west shore south of ')
add_hyperlink(p, 'Dundee', 'https://www.google.com/maps/search/?api=1&query=Dundee+NY', MAPS_GREEN)
add_run(p, '. The 30-room inn sits above the vineyard, each room facing the lake from a private balcony or patio. Veraisons Restaurant changes its menu each season, sourcing from local farms throughout the Finger Lakes, with an unobstructed view of the vineyard and water below. Open year-round for tastings, dining, and lodging. Ten minutes north of ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, ' on Route 14.')

print("Part 4 done: wine trail + tier 1")

# Tier 2 producers
add_heading_styled('Tier 2', level=3)

tier2 = [
    ('Boundary Breaks Vineyard', 'One vineyard, one variety, no buses',
     'Bruce Murray purchased 120 acres of bare land near Lodi Point in 2008 and planted vines. The focus has been single-vineyard Riesling from the start: precise, mineral-driven wines from shale slopes close to the water. Wine Enthusiast has named the Boundary Breaks Dry Riesling to its Top 100 Wines in the World twice. No buses, no limos, no groups larger than six. The tasting room sits just north of Lodi Point State Park with an unobstructed lake view. Open daily 11am to 5pm year-round.',
     'https://www.google.com/maps/search/?api=1&query=Boundary+Breaks+Vineyard+Lodi+NY'),
    ('Hillick & Hobbs Estate', 'Small production, longer conversations',
     'Hillick & Hobbs Estate sits on Route 79 in Burdett at the south end of the east shore, making it the first purpose-built small-production stop on the Route 414 corridor. The estate grows Riesling on a single vineyard above the lake, farming the shale slopes that give east-shore Seneca its claim to Mosel comparison. Production is deliberately small. Fewer visitors, longer pours, more room to ask questions about the vineyard. Close enough to the Burdett cluster that the east-shore entry point has a full half-day of stops before you even leave Burdett.',
     'https://www.google.com/maps/search/?api=1&query=Hillick+and+Hobbs+Estate+Burdett+NY'),
    ('Silver Thread Vineyard', 'Organic since 1994, before it was a selling point',
     'Silver Thread Vineyard has been certified organic since 1994, a commitment made long before organic viticulture became a marketing category in the Finger Lakes. The east-shore estate outside Lodi farms the shale slopes with no synthetic inputs and makes wine with minimal intervention in the cellar. The Rieslings are precise and restrained, built for people who want to taste the hillside rather than the winemaker\'s hand. Production is small; capacity in the tasting room is limited.',
     'https://www.google.com/maps/search/?api=1&query=Silver+Thread+Vineyard+Lodi+NY'),
    ('Hazlitt 1852 Vineyards', 'Farming this hillside since before the Civil War',
     'The Hazlitt family has grown grapes on this east-shore hillside in Hector since 1852, which puts them among the longest-running agricultural families in the region. The winery produces a wide portfolio, from estate Riesling and Cabernet Franc to Red Cat, one of the best-selling wines in New York State. The Cider Tree hard cider operation runs alongside on the same property, earning both a wine trail and craft beverage trail badge on a single visit.',
     'https://www.google.com/maps/search/?api=1&query=Hazlitt+1852+Vineyards+Hector+NY'),
    ('Lamoreaux Landing Wine Cellars', 'Three generations, Greek Revival, east-shore Riesling',
     'The Lamoreaux family has been farming grapes outside Lodi for three generations. The Greek Revival tasting room looks directly down the hillside to the lake, one of the cleaner views from any tasting room on the east-shore drive. The estate produces Riesling, Chardonnay, Cabernet Franc, and Pinot Noir from shale-slope vineyards above the water. A consistent, serious program with a straightforward tasting experience.',
     'https://www.google.com/maps/search/?api=1&query=Lamoreaux+Landing+Wine+Cellars+Lodi+NY'),
    ('Atwater Vineyards', 'First winery north of Watkins Glen on Route 414',
     'Atwater Vineyards occupies the east-shore bluff above Seneca Lake in Burdett, the first substantial winery stop heading north on Route 414 from Watkins Glen. The estate produces Riesling, Gewurztraminer, and red varieties from vineyards with direct lake exposure. An approachable operation with enough range to orient visitors still figuring out what east-shore Seneca wine actually tastes like.',
     'https://www.google.com/maps/search/?api=1&query=Atwater+Vineyards+Burdett+NY'),
    ('Ravines Wine Cellars', 'Burgundy-trained Riesling, Geneva tasting room',
     'Morten Hallgren trained at Domaine de la Romanee-Conti in Burgundy before founding Ravines on the east shore of Seneca Lake. The estate farming focuses on Riesling and Cabernet Franc with an approach shaped by Old World technique. The Geneva tasting room on Barracks Road is the practical access point for visitors based at the north end of the lake. A second location in Hammondsport covers Keuka Lake for visitors running the full Finger Lakes circuit.',
     'https://www.google.com/maps/search/?api=1&query=Ravines+Wine+Cellars+Geneva+NY'),
]

for name, subtitle, desc, maps_url in tier2:
    t2h = doc.add_paragraph()
    add_hyperlink(t2h, name, maps_url, MAPS_GREEN)
    sub = doc.add_paragraph()
    add_run(sub, subtitle, italic=True)
    p = doc.add_paragraph()
    add_run(p, desc)

# ============================================================
# FARMS, CIDER, AND FOOD (split from combined section)
# ============================================================
add_heading_styled('Farms, cider, and food', level=2)

# --- Hazlitt / Cider Tree ---
farm_p = doc.add_paragraph()
add_run(farm_p, '\u2014 ')
add_hyperlink(farm_p, 'Hazlitt 1852 Vineyards / The Cider Tree', 'https://www.google.com/maps/search/?api=1&query=Hazlitt+1852+Vineyards+Hector+NY', MAPS_GREEN)
add_run(farm_p, ' ')
add_hyperlink(farm_p, 'Hector', 'https://www.google.com/maps/search/?api=1&query=Hector+NY', MAPS_GREEN)
sub_p = doc.add_paragraph()
add_run(sub_p, 'Two trail badges on one stop since before the Civil War', italic=True)
doc.add_paragraph('The Hazlitt family has farmed this east-shore hillside since 1852. The Cider Tree hard cider operation runs alongside the winery on the same property, earning both a wine trail and Craft Beverage Trail badge on a single visit. The breadth is intentional: a working farm operation that serves everyone from serious collectors to first-time trail visitors.')

# --- Two Goats Brewing ---
farm_p = doc.add_paragraph()
add_run(farm_p, '\u2014 ')
add_hyperlink(farm_p, 'Two Goats Brewing', 'https://www.google.com/maps/search/?api=1&query=Two+Goats+Brewing+Burdett+NY', MAPS_GREEN)
add_run(farm_p, ' ')
add_hyperlink(farm_p, 'Burdett', 'https://www.google.com/maps/search/?api=1&query=Burdett+NY', MAPS_GREEN)
sub_p = doc.add_paragraph()
add_run(sub_p, 'The east-shore stop for people who don\'t drink wine', italic=True)
doc.add_paragraph('Craft brewery in the Burdett cluster on Route 414. Close enough to Hillick & Hobbs and Finger Lakes Distilling that the first hour on the east shore covers wine, beer, and spirits without moving the car far. Craft Beverage Trail badge.')

# --- Finger Lakes Distilling ---
farm_p = doc.add_paragraph()
add_run(farm_p, '\u2014 ')
add_hyperlink(farm_p, 'Finger Lakes Distilling', 'https://www.google.com/maps/search/?api=1&query=Finger+Lakes+Distilling+Burdett+NY', MAPS_GREEN)
add_run(farm_p, ' ')
add_hyperlink(farm_p, 'Burdett', 'https://www.google.com/maps/search/?api=1&query=Burdett+NY', MAPS_GREEN)
sub_p = doc.add_paragraph()
add_run(sub_p, 'Grain-to-glass spirits on the east-shore bluff', italic=True)
doc.add_paragraph('Craft distillery producing gin, vodka, whiskey, and grappa from local grain and grape. Part of the Burdett cluster that makes the east-shore entry point the most complete first stop on the trail. Craft Beverage Trail badge.')

# --- Blueberry Patch ---
farm_p = doc.add_paragraph()
add_run(farm_p, '\u2014 Blueberry Patch ')
add_hyperlink(farm_p, 'Finger Lakes National Forest', 'https://www.google.com/maps/search/?api=1&query=Finger+Lakes+National+Forest+Hector+NY', MAPS_GREEN)
sub_p = doc.add_paragraph()
add_run(sub_p, 'U-pick blueberries in a national forest', italic=True)
doc.add_paragraph('Adjacent to the nine-site first-come campground in the Finger Lakes National Forest. U-pick blueberries in season. The kind of stop that doesn\'t appear on wine trail itineraries, which is part of its appeal.')

print("Part 5 done: tier 2 and farms")

# ============================================================
# OUTDOOR AND NATURE (split from combined section)
# ============================================================
add_heading_styled('Outdoor and nature', level=2)

p = doc.add_paragraph()
add_run(p, 'The gorge at ')
add_hyperlink(p, 'Watkins Glen State Park', 'https://en.wikipedia.org/wiki/Watkins_Glen_State_Park', WIKI_BLUE)
add_run(p, ' is the primary natural draw on Seneca. Nineteen waterfalls in 1.5 miles, 800 stone steps, canyon walls that the Civilian Conservation Corps rebuilt in natural stone after a 1935 flood. Arrive before 9am in summer.')

p = doc.add_paragraph()
add_run(p, 'Seneca also has ')
add_hyperlink(p, 'Shequaga Falls', 'https://www.google.com/maps/search/?api=1&query=Shequaga+Falls+Montour+Falls+NY', MAPS_GREEN)
add_run(p, ' in the center of ')
add_hyperlink(p, 'Montour Falls', 'https://www.google.com/maps/search/?api=1&query=Montour+Falls+NY', MAPS_GREEN)
add_run(p, ' village, five minutes south of ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, '. One hundred sixty-five feet. Most people drive through without stopping. That is a mistake.')

p = doc.add_paragraph()
add_run(p, 'The ')
add_hyperlink(p, 'Finger Lakes National Forest', 'https://www.google.com/maps/search/?api=1&query=Finger+Lakes+National+Forest+Hector+NY', MAPS_GREEN)
add_run(p, ' is New York\'s only national forest, 16,000 acres on the ridge between Seneca and Cayuga Lakes in Schuyler and Seneca Counties. Between 1938 and 1941, over 100 Depression-era farms were purchased by the federal government after soil depletion made them unviable. The Forest Service still leases the pasture land to working farms: hikers on the 30-mile trail system may encounter free-ranging cattle alongside vineyard views. The Interloken Trail runs 12 miles north-south through the forest. The Blueberry Patch campground is nine sites, first-come, with a U-pick blueberry patch adjacent. This is the east-shore outdoor stop that does not appear on most wine trail itineraries.')

# Named sites list
add_heading_styled('Natural sites', level=3)
nat_sites = [
    ('Watkins Glen State Park', 'Watkins Glen NY', 'The best gorge trail in the northeast', '19 waterfalls in 1.5 miles. The gorge trail winds through tunnels, over and under falls, past 200-foot water-carved shale walls. 800 stone steps, rebuilt by the CCC after a 1935 flood. Gorge trail open mid-May through late October; rim trails year-round. Get there before 9am in summer or after 4pm. The gorge fills up.', 'https://en.wikipedia.org/wiki/Watkins_Glen_State_Park', WIKI_BLUE),
    ('Shequaga Falls', 'Montour Falls NY', 'Most people drive through without stopping. That is a mistake.', '165-foot waterfall in the center of Montour Falls village, five minutes south of Watkins Glen. The falls are visible from the main road. There is no trail required, no fee, and no crowd. One of the most undervisited natural features in the Finger Lakes.', 'https://www.google.com/maps/search/?api=1&query=Shequaga+Falls+Montour+Falls+NY', MAPS_GREEN),
    ('Hector Falls', 'Burdett NY', 'A pull-over waterfall on the east-shore drive', 'Visible from Route 414 on the east shore. The upper cascades are seen from the road; the lower cascades are visible from the lake. Pull over safely. Worth combining with the first east-shore winery stops at Atwater and Hillick & Hobbs nearby.', 'https://www.google.com/maps/search/?api=1&query=Hector+Falls+Burdett+NY', MAPS_GREEN),
    ('Finger Lakes National Forest', 'Hector NY', 'New York\'s only national forest, and the cattle have the right of way', 'Sixteen thousand acres on the ridge between Seneca and Cayuga Lakes. Over 100 Depression-era farms were purchased by the federal government between 1938 and 1941. The Forest Service still leases the pastures to working farms: hikers on the 30-mile trail system may encounter free-ranging cattle alongside vineyard views. The Interloken Trail runs 12 miles north-south. Blueberry Patch campground is nine first-come sites with a U-pick patch adjacent.', 'https://www.google.com/maps/search/?api=1&query=Finger+Lakes+National+Forest+Hector+NY', MAPS_GREEN),
    ('Seneca Lake State Park', 'Geneva NY', 'Known for the free public beach at the north end', 'Free public beach, boat launch, and marina at the lake\'s north end in Geneva. The practical starting point for visitors based in Geneva who want lake access without a winery attached. Picnic facilities, playground, and Sprayground water feature in summer.', 'https://www.google.com/maps/search/?api=1&query=Seneca+Lake+State+Park+Geneva+NY', MAPS_GREEN),
]
for name, loc, subtitle, desc, url, color in nat_sites:
    p = doc.add_paragraph()
    add_run(p, '\u2014 ')
    add_hyperlink(p, name, url, color)
    add_run(p, ' ')
    add_hyperlink(p, loc, f'https://www.google.com/maps/search/?api=1&query={loc.replace(" ", "+").replace(",", "")}', MAPS_GREEN)
    sub_p = doc.add_paragraph()
    add_run(sub_p, subtitle, italic=True)
    doc.add_paragraph(desc)

# ============================================================
# GETTING HERE AND GETTING AROUND (moved to canonical position)
# ============================================================
add_heading_styled('Getting here and getting around', level=2)

p = doc.add_paragraph()
add_run(p, 'Seneca Lake requires a car. ')
add_hyperlink(p, 'Syracuse Hancock International Airport', 'https://www.google.com/maps/search/?api=1&query=Syracuse+Hancock+International+Airport', MAPS_GREEN)
add_run(p, ' is about an hour northeast. ')
add_hyperlink(p, 'Greater Rochester International Airport', 'https://www.google.com/maps/search/?api=1&query=Greater+Rochester+International+Airport', MAPS_GREEN)
add_run(p, ' is about an hour northwest. From New York City: four hours, I-86 to Route 14.')

# Shore roads
add_heading_styled('The two shore roads', level=3)
p = doc.add_paragraph()
add_run(p, 'Route 14 runs the west shore from ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, ' north through ')
add_hyperlink(p, 'Rock Stream', 'https://www.google.com/maps/search/?api=1&query=Rock+Stream+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Dundee', 'https://www.google.com/maps/search/?api=1&query=Dundee+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Himrod', 'https://www.google.com/maps/search/?api=1&query=Himrod+NY', MAPS_GREEN)
add_run(p, ', to ')
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, '. More producers, more traffic. Route 414 runs the east shore north through ')
add_hyperlink(p, 'Burdett', 'https://www.google.com/maps/search/?api=1&query=Burdett+NY', MAPS_GREEN)
add_run(p, ', ')
add_hyperlink(p, 'Hector', 'https://www.google.com/maps/search/?api=1&query=Hector+NY', MAPS_GREEN)
add_run(p, ', and ')
add_hyperlink(p, 'Lodi', 'https://www.google.com/maps/search/?api=1&query=Lodi+NY', MAPS_GREEN)
add_run(p, '. Quieter, steeper, the better second day. The two roads connect at ')
add_hyperlink(p, 'Geneva', 'https://www.google.com/maps/search/?api=1&query=Geneva+NY', MAPS_GREEN)
add_run(p, ' on the north end and ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+NY', MAPS_GREEN)
add_run(p, ' on the south.')

# Drive times
add_heading_styled('Drive times', level=3)
drive_data = [
    ('Origin', 'Destination', 'Time'),
    ('Watkins Glen', 'Geneva', '50 minutes'),
    ('Watkins Glen', 'Lodi', '20 minutes'),
    ('Geneva', 'Seneca Falls', '20 minutes'),
    ('Geneva', 'Ithaca', '45 minutes'),
    ('Geneva', 'Auburn', '35 minutes'),
]
drive_table = doc.add_table(rows=len(drive_data), cols=3)
drive_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (origin, dest, time) in enumerate(drive_data):
    row = drive_table.rows[i]
    for j, val in enumerate([origin, dest, time]):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

# Seasonal notes
add_heading_styled('Seasonal notes', level=3)
p = doc.add_paragraph()
add_run(p, 'Wine trail tasting rooms: full hours May through November, reduced or appointment-only December through April, call ahead. Gorge trail at ')
add_hyperlink(p, 'Watkins Glen', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+State+Park+NY', MAPS_GREEN)
add_run(p, ': mid-May through late October. Rim trails: year-round. ')
add_hyperlink(p, 'Finger Lakes National Forest', 'https://www.google.com/maps/search/?api=1&query=Finger+Lakes+National+Forest+Hector+NY', MAPS_GREEN)
add_run(p, ' trails: year-round.')

# ============================================================
# STAMP AND BADGE CALLOUT
# ============================================================
add_heading_styled('Stamp and badge callout', level=2)

wine_p = doc.add_paragraph()
add_run(wine_p, 'Wine trail stamps: ', bold=True)
add_run(wine_p, '33 producers tagged to the ')
add_hyperlink(wine_p, 'Seneca Lake Wine Trail', 'https://upstate.tourismo.app/trails/seneca-lake-wine-trail', UPSTATE_RED)
add_run(wine_p, ' badge. Count toward the Finger Lakes Wine master badge.')

heritage_p = doc.add_paragraph()
add_run(heritage_p, 'Heritage stamps: ', bold=True)
h_sites = ["Women's Rights NHP", "M'Clintock House", 'National Memorial Day Museum', 'Geneva Experiment Station']
add_run(heritage_p, ', '.join(h_sites))
add_run(heritage_p, ' \u2014 all ')
add_hyperlink(heritage_p, 'Finger Lakes Heritage Trail', 'https://upstate.tourismo.app/trails/finger-lakes-heritage-trail', UPSTATE_RED)
add_run(heritage_p, ' badge.')

nature_p = doc.add_paragraph()
add_run(nature_p, 'Nature stamps: ', bold=True)
add_hyperlink(nature_p, 'Watkins Glen State Park', 'https://www.google.com/maps/search/?api=1&query=Watkins+Glen+State+Park+NY', MAPS_GREEN)
add_run(nature_p, ', ')
add_hyperlink(nature_p, 'Shequaga Falls', 'https://www.google.com/maps/search/?api=1&query=Shequaga+Falls+Montour+Falls+NY', MAPS_GREEN)
add_run(nature_p, ', ')
add_hyperlink(nature_p, 'Finger Lakes National Forest', 'https://www.google.com/maps/search/?api=1&query=Finger+Lakes+National+Forest+Hector+NY', MAPS_GREEN)
add_run(nature_p, ', ')
add_hyperlink(nature_p, 'Seneca Lake State Park', 'https://www.google.com/maps/search/?api=1&query=Seneca+Lake+State+Park+Geneva+NY', MAPS_GREEN)
add_run(nature_p, ' \u2014 Finger Lakes Outdoor badge.')

# ============================================================
# CTA BLOCK
# ============================================================
cta_heading = doc.add_paragraph()
add_run(cta_heading, 'Explore more of the Finger Lakes.', bold=True, size=13)
apply_cta_shading(cta_heading)

cta_p1 = doc.add_paragraph()
add_run(cta_p1, 'Seneca is the wine center. The ')
add_hyperlink(cta_p1, 'Keuka Lake chapter', 'https://upstate.tourismo.app/itineraries/keuka-lake', UPSTATE_RED)
add_run(cta_p1, ' is where the wine story actually started, twenty miles west on a Y-shaped lake above Hammondsport. The ')
add_hyperlink(cta_p1, 'Cayuga Lake chapter', 'https://upstate.tourismo.app/itineraries/cayuga-lake', UPSTATE_RED)
add_run(cta_p1, ' picks up the reform corridor. The ')
add_hyperlink(cta_p1, 'Finger Lakes regional guide', 'https://upstatebound.com/guides/finger-lakes-region-04b4ec42-84de-4c38-aa0a-f689dc88d7a6', UPSTATE_RED)
add_run(cta_p1, ' connects all five.')
apply_cta_shading(cta_p1)

cta_p2 = doc.add_paragraph()
add_hyperlink(cta_p2, 'Plan your visit on Upstate \u2192', 'https://upstate.tourismo.app/itineraries/seneca-lake', UPSTATE_RED)
apply_cta_shading(cta_p2)

cta_p3 = doc.add_paragraph()
add_run(cta_p3, 'Related: ')
add_hyperlink(cta_p3, 'Around Seneca Wine Itinerary', 'https://upstate.tourismo.app/itineraries/around-seneca', UPSTATE_RED)
apply_cta_shading(cta_p3)

# ============================================================
# SAVE
# ============================================================
output_dir = '/mnt/user-data/outputs'
os.makedirs(output_dir, exist_ok=True)

output_path = os.path.join(output_dir, 'Seneca_Lake_Chapter_v2.docx')
doc.save(output_path)
print(f"\nSaved to: {output_path}")

repo_path = '/home/user/fingerlakes/Seneca_Lake_Chapter_v2.docx'
doc.save(repo_path)
print(f"Also saved to: {repo_path}")
print("\nDone! Seneca chapter includes:")
print("  - TLDR added")
print("  - Intro: 4 paragraphs (physical, Indigenous, wine/ag, connecting)")
print("  - Section order fixed to canonical template")
print("  - Farms and Outdoor split into separate sections")
print("  - Producer count fixed to 33 throughout")
print("  - Heritage Corridor renamed")
print("  - Color-coded hyperlinks throughout")
print("  - Stamp/badge callout reformatted")
print("  - CTA block added")
print("  - BEST FOR subtitles on all 6 towns")
